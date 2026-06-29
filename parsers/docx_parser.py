"""
Parser de documentos DOCX.

Diferente do PDF, o formato DOCX possui um conceito NATIVO de tabela
(`document.tables`), o que torna a extração muito mais confiável do que
a heurística usada no parser de PDF — não precisamos "adivinhar" onde
uma tabela começa e termina.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from parsers.raw_structures import RawDocument, RawTable, RawTableCell, RawTextLine
from utils.logger import get_logger

logger = get_logger(__name__)


class DOCXParser:
    """Extrai texto e tabelas de arquivos DOCX (Microsoft Word)."""

    def parse(self, file_path: str | Path) -> RawDocument:
        """Extrai todo o conteúdo textual e tabular de um DOCX.

        Percorre o documento na ordem de leitura real (parágrafos e
        tabelas intercalados como aparecem no arquivo), preservando a
        ordem visual original — importante para que a fase de
        Classificação consiga inferir corretamente, por exemplo, que um
        título "Avaliação de Bioimpedância" antecede a tabela que o segue.

        Args:
            file_path: Caminho do arquivo .docx no disco.

        Returns:
            `RawDocument` com linhas de texto e tabelas detectadas.
        """
        file_path = Path(file_path)
        logger.info("Iniciando parsing de DOCX: %s", file_path.name)

        docx_doc = DocxDocument(str(file_path))
        document = RawDocument(nome_arquivo=file_path.name)

        ordem_global = 0
        for elemento in self._iter_block_items(docx_doc):
            if elemento["tipo"] == "paragrafo":
                texto = elemento["texto"].strip()
                if texto:
                    document.linhas.append(
                        RawTextLine(texto=texto, ordem=ordem_global, pagina=0)
                    )
                    ordem_global += 1
            else:  # tabela
                tabela = self._extract_table(elemento["tabela"], ordem_global)
                document.tabelas.append(tabela)
                ordem_global += 1

        logger.info(
            "DOCX parseado: %d linhas de texto, %d tabelas detectadas.",
            len(document.linhas), len(document.tabelas),
        )
        return document

    def _iter_block_items(self, docx_doc) -> list[dict]:
        """Itera parágrafos e tabelas na ordem real do documento.

        A API pública do python-docx expõe `document.paragraphs` e
        `document.tables` como listas SEPARADAS, perdendo a ordem
        relativa entre eles. Para preservar a ordem visual, percorremos
        diretamente a árvore XML do corpo do documento.
        """
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        itens = []
        body = docx_doc.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                paragrafo = Paragraph(child, docx_doc)
                itens.append({"tipo": "paragrafo", "texto": paragrafo.text})
            elif child.tag == qn("w:tbl"):
                tabela = Table(child, docx_doc)
                itens.append({"tipo": "tabela", "tabela": tabela})
        return itens

    def _extract_table(self, docx_table, ordem: int) -> RawTable:
        """Converte uma tabela nativa do python-docx em `RawTable`."""
        num_colunas = len(docx_table.columns)
        raw_table = RawTable(num_linhas=len(docx_table.rows), num_colunas=num_colunas, ordem=ordem)

        for linha_index, linha in enumerate(docx_table.rows):
            for coluna_index, celula in enumerate(linha.cells):
                raw_table.celulas.append(
                    RawTableCell(texto=celula.text.strip(), linha=linha_index, coluna=coluna_index)
                )

        return raw_table
