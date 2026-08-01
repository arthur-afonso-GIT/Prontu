"""Regras puras dos modelos e formulários de fichas clínicas."""
from __future__ import annotations

import html
import re
import unicodedata
from datetime import date, datetime


NOME_MODELO_PADRAO = "Ficha de Consulta Geral (Padrão)"

MODELO_PADRAO = [
    {"tipo": "secao", "label": "HISTÓRICO DA CONSULTA ATUAL"},
    {
        "tipo": "texto_longo",
        "label": "Queixa Principal (QP)",
        "id": "qp",
    },
    {
        "tipo": "texto_longo",
        "label": "Histórico da Doença Atual (HDA)",
        "id": "hda",
    },
    {"tipo": "secao", "label": "EXAME FÍSICO E SINAIS VITAIS"},
    {
        "tipo": "texto_curto",
        "label": "Pressão Arterial (PA)",
        "id": "pa",
        "placeholder": "Ex.: 120x80 mmHg",
    },
    {
        "tipo": "texto_curto",
        "label": "Frequência Cardíaca (FC)",
        "id": "fc",
        "placeholder": "Ex.: 75 bpm",
    },
    {"tipo": "secao", "label": "CONDUTA MÉDICA"},
    {
        "tipo": "texto_longo",
        "label": "Prescrição / Orientações Passadas",
        "id": "prescricao",
    },
]

TIPOS_SUPORTADOS = {
    "secao",
    "texto_curto",
    "texto_longo",
    "checkbox",
    "numero",
    "data",
    "multipla_escolha",
}

SEMANTICAS_SUPORTADAS = {
    "data",
    "data_nascimento",
    "idade",
    "telefone",
    "cpf",
    "rg",
}


def _slug(texto: str) -> str:
    normalizado = unicodedata.normalize("NFKD", str(texto or "campo"))
    sem_acentos = "".join(
        caractere for caractere in normalizado
        if not unicodedata.combining(caractere)
    )
    return re.sub(r"[^a-z0-9]+", "_", sem_acentos.casefold()).strip("_") or "campo"


def rotulo_legivel(campo: dict) -> str:
    """Remove identificadores antigos que chegaram a aparecer para o usuário."""
    label = " ".join(str(campo.get("label") or "").split())
    campo_id = str(campo.get("id") or "")
    partes = campo_id.split("_")
    if len(partes) >= 3 and partes[0].casefold() == "custom":
        identificador = partes[1]
        if label.casefold().startswith(f"{identificador.casefold()} "):
            label = label[len(identificador):].strip()
    if label:
        return label
    partes_id = [
        parte for parte in partes
        if parte and parte.casefold() != "custom"
        and not re.fullmatch(r"[0-9a-f]{8}", parte, re.IGNORECASE)
    ]
    return " ".join(partes_id).replace("_", " ").strip().capitalize() or "Campo"


def _chave_semantica(texto: object) -> str:
    normalizado = unicodedata.normalize("NFKD", str(texto or ""))
    sem_acentos = "".join(
        caractere for caractere in normalizado
        if not unicodedata.combining(caractere)
    )
    return re.sub(r"[^a-z0-9]+", " ", sem_acentos.casefold()).strip()


def _inferir_semantica(campo: dict) -> str:
    """Reconhece o significado clínico sem depender do nome exato do campo."""
    explicacao = " ".join(
        str(campo.get(chave) or "")
        for chave in ("label", "id", "ajuda", "placeholder")
    )
    chave = _chave_semantica(explicacao)
    tokens = set(chave.split())

    if "cpf" in tokens:
        return "cpf"
    if "rg" in tokens or "identidade" in tokens:
        return "rg"
    if tokens.intersection({"telefone", "celular", "whatsapp", "fone"}):
        return "telefone"
    if "idade" in tokens or "anos" in tokens and campo.get("unidade") == "anos":
        return "idade"
    if (
        "nascimento" in tokens
        and tokens.intersection({"data", "nascimento", "nasc"})
    ) or "data nascimento" in chave or "data de nascimento" in chave:
        return "data_nascimento"
    if campo.get("tipo") == "data" or "data" in tokens:
        return "data"
    return ""


def _relacionar_campos_calculados(campos: list[dict]) -> None:
    """Liga um campo de idade à data de nascimento mais coerente."""
    grupo = 0
    grupos: dict[int, int] = {}
    campos_dados: list[tuple[int, dict]] = []
    for indice, campo in enumerate(campos):
        if campo.get("tipo") == "secao":
            grupo += 1
            continue
        grupos[indice] = grupo
        campos_dados.append((indice, campo))

    datas_nascimento = [
        (indice, campo)
        for indice, campo in campos_dados
        if campo.get("semantica") == "data_nascimento"
    ]
    datas_genericas = [
        (indice, campo)
        for indice, campo in campos_dados
        if campo.get("semantica") == "data"
    ]

    for indice_idade, campo_idade in campos_dados:
        if campo_idade.get("semantica") != "idade":
            continue
        origem_existente = str(campo_idade.get("calculado_por") or "").strip()
        if origem_existente:
            campo_idade["somente_leitura"] = True
            continue

        mesmo_grupo = [
            item for item in datas_nascimento
            if grupos.get(item[0]) == grupos.get(indice_idade)
        ]
        candidatas = mesmo_grupo or datas_nascimento

        if not candidatas:
            # "Data" seguido de "Idade" é um padrão comum em fichas importadas.
            vizinhas = [
                item for item in datas_genericas
                if grupos.get(item[0]) == grupos.get(indice_idade)
                and abs(item[0] - indice_idade) <= 2
            ]
            candidatas = vizinhas

        if not candidatas:
            continue
        _, origem = min(
            candidatas,
            key=lambda item: (abs(item[0] - indice_idade), item[0] > indice_idade),
        )
        origem["semantica"] = "data_nascimento"
        campo_idade["calculado_por"] = origem["id"]
        campo_idade["somente_leitura"] = True
        campo_idade.setdefault("unidade", "anos")


def normalizar_estrutura(estrutura) -> list[dict]:
    """Entrega à interface somente campos válidos, estáveis e serializáveis."""
    if not isinstance(estrutura, list):
        return []
    normalizados = []
    ids_usados: set[str] = set()
    for indice, original in enumerate(estrutura):
        if not isinstance(original, dict):
            continue
        tipo = str(original.get("tipo") or "").strip()
        if tipo not in TIPOS_SUPORTADOS:
            continue
        campo = dict(original)
        campo["tipo"] = tipo
        campo["label"] = rotulo_legivel(campo)
        if tipo != "secao":
            campo_id = str(campo.get("id") or "").strip()
            if not campo_id:
                campo_id = f"campo_{indice}_{_slug(campo['label'])}"
            base = campo_id
            contador = 2
            while campo_id in ids_usados:
                campo_id = f"{base}_{contador}"
                contador += 1
            campo["id"] = campo_id
            ids_usados.add(campo_id)
            campo["obrigatorio"] = bool(campo.get("obrigatorio", False))
            semantica = str(campo.get("semantica") or "").strip()
            if semantica not in SEMANTICAS_SUPORTADAS:
                semantica = _inferir_semantica(campo)
            if semantica:
                campo["semantica"] = semantica
        if tipo == "multipla_escolha":
            campo["opcoes"] = [
                str(opcao).strip()
                for opcao in campo.get("opcoes", [])
                if str(opcao).strip()
            ]
        normalizados.append(campo)
    _relacionar_campos_calculados(normalizados)
    return normalizados


def _aplicar_mascara(digitos: str, grupos: tuple[int, ...], separadores: tuple[str, ...]) -> str:
    partes = []
    inicio = 0
    for tamanho in grupos:
        trecho = digitos[inicio:inicio + tamanho]
        if not trecho:
            break
        partes.append(trecho)
        inicio += tamanho
    resultado = partes[0] if partes else ""
    for indice, parte in enumerate(partes[1:]):
        resultado += separadores[indice] + parte
    return resultado


def formatar_data_ficha(valor: object) -> str:
    digitos = re.sub(r"\D", "", str(valor or ""))[:8]
    return _aplicar_mascara(digitos, (2, 2, 4), ("/", "/"))


def formatar_telefone_ficha(valor: object) -> str:
    digitos = re.sub(r"\D", "", str(valor or ""))[:11]
    if not digitos:
        return ""
    if len(digitos) <= 2:
        return f"({digitos}"
    tamanho_prefixo = 5 if len(digitos) > 10 else 4
    prefixo = digitos[2:2 + tamanho_prefixo]
    sufixo = digitos[2 + tamanho_prefixo:]
    resultado = f"({digitos[:2]}) {prefixo}"
    return resultado + (f"-{sufixo}" if sufixo else "")


def formatar_cpf_ficha(valor: object) -> str:
    digitos = re.sub(r"\D", "", str(valor or ""))[:11]
    return _aplicar_mascara(digitos, (3, 3, 3, 2), (".", ".", "-"))


def formatar_rg_ficha(valor: object) -> str:
    texto = re.sub(r"[^0-9A-Za-z]", "", str(valor or "")).upper()[:14]
    if texto.isdigit() and len(texto) <= 9:
        return _aplicar_mascara(texto, (2, 3, 3, 1), (".", ".", "-"))
    return texto


def normalizar_valor_campo(campo: dict, valor: object) -> object:
    """Aplica a apresentação adequada ao significado do campo."""
    semantica = str((campo or {}).get("semantica") or "")
    if semantica in {"data", "data_nascimento"}:
        return formatar_data_ficha(valor)
    if semantica == "telefone":
        return formatar_telefone_ficha(valor)
    if semantica == "cpf":
        return formatar_cpf_ficha(valor)
    if semantica == "rg":
        return formatar_rg_ficha(valor)
    return valor


def calcular_idade(data_nascimento: object, hoje: date | None = None) -> str:
    try:
        nascimento = datetime.strptime(
            formatar_data_ficha(data_nascimento), "%d/%m/%Y"
        ).date()
    except (TypeError, ValueError):
        return ""
    referencia = hoje or date.today()
    if nascimento > referencia:
        return ""
    idade = referencia.year - nascimento.year - (
        (referencia.month, referencia.day)
        < (nascimento.month, nascimento.day)
    )
    return str(idade) if 0 <= idade <= 130 else ""


def atualizar_respostas_calculadas(
    campos: list[dict],
    respostas: dict,
    hoje: date | None = None,
) -> dict:
    """Normaliza respostas e recalcula campos derivados, como a idade."""
    campos = normalizar_estrutura(campos)
    resultado = dict(respostas or {})
    por_id = {
        str(campo.get("id")): campo
        for campo in campos
        if campo.get("id")
    }
    for campo_id, campo in por_id.items():
        if campo_id in resultado and campo.get("semantica") != "idade":
            resultado[campo_id] = normalizar_valor_campo(
                campo, resultado[campo_id]
            )
    for campo_id, campo in por_id.items():
        if campo.get("semantica") != "idade":
            continue
        origem = str(campo.get("calculado_por") or "")
        if origem:
            resultado[campo_id] = calcular_idade(resultado.get(origem), hoje)
    return resultado


def respostas_iniciais(campos: list[dict]) -> dict:
    hoje = datetime.now().strftime("%d/%m/%Y")
    respostas = {}
    for campo in campos:
        campo_id = campo.get("id")
        if not campo_id:
            continue
        if campo.get("tipo") == "checkbox":
            respostas[campo_id] = False
        elif campo.get("tipo") == "data" and campo.get("preencher_hoje"):
            respostas[campo_id] = hoje
        else:
            respostas[campo_id] = ""
    return respostas


def validar_respostas(
    campos: list[dict], respostas: dict
) -> tuple[bool, list[str]]:
    vazios = []
    for campo in campos:
        if campo.get("tipo") == "secao" or not campo.get("obrigatorio"):
            continue
        valor = respostas.get(campo.get("id"))
        if valor is None or valor is False or not str(valor).strip():
            vazios.append(campo.get("label") or "Campo")
    return not vazios, vazios


def preparar_dados_exportacao(
    campos: list[dict],
    respostas: dict,
    paciente: str,
    modelo: str,
    clinica: str = "",
    profissional: str = "",
    data: str | None = None,
) -> dict:
    """Transforma o formulário dinâmico em dados estáveis para Word e PDF."""
    itens = []
    respostas = dict(respostas or {})
    for campo in normalizar_estrutura(campos):
        rotulo = rotulo_legivel(campo)
        if campo.get("tipo") == "secao":
            itens.append({"tipo": "secao", "rotulo": rotulo, "valor": ""})
            continue
        valor = respostas.get(campo.get("id"))
        if campo.get("tipo") == "checkbox":
            valor = "Sim" if bool(valor) else "Não"
        elif isinstance(valor, (list, tuple)):
            valor = ", ".join(str(item) for item in valor if str(item).strip())
        else:
            valor = str(valor or "").strip()
        itens.append({
            "tipo": "campo",
            "rotulo": rotulo,
            "valor": valor or "Não informado",
        })
    return {
        "paciente": str(paciente or "").strip(),
        "modelo": str(modelo or "Ficha Clínica").strip(),
        "clinica": str(clinica or "").strip(),
        "profissional": str(profissional or "").strip(),
        "data": data or datetime.now().strftime("%d/%m/%Y às %H:%M"),
        "itens": itens,
    }


def nome_arquivo_exportacao(dados: dict, extensao: str) -> str:
    nome = unicodedata.normalize("NFKD", str(dados.get("paciente") or "Paciente"))
    nome = "".join(
        caractere for caractere in nome
        if not unicodedata.combining(caractere)
    )
    nome = re.sub(r"[^A-Za-z0-9_-]+", "_", nome).strip("_") or "Paciente"
    extensao = extensao if extensao.startswith(".") else f".{extensao}"
    return f"Ficha_{nome}_{datetime.now().strftime('%Y%m%d_%H%M')}{extensao}"


def exportar_ficha_word(
    dados: dict,
    caminho: str,
) -> None:
    """Gera um documento Word clínico, sem marcas do sistema."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    documento = Document()
    secao = documento.sections[0]
    secao.top_margin = Inches(0.65)
    secao.bottom_margin = Inches(0.65)

    titulo = documento.add_heading(dados.get("modelo") or "Ficha Clínica", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    documento.add_paragraph(
        f"Paciente: {dados.get('paciente') or 'Não informado'}"
    )
    documento.add_paragraph(f"Data: {dados.get('data') or ''}")
    if dados.get("clinica"):
        documento.add_paragraph(f"Clínica: {dados['clinica']}")
    if dados.get("profissional"):
        documento.add_paragraph(f"Profissional: {dados['profissional']}")

    for item in dados.get("itens") or []:
        if item.get("tipo") == "secao":
            documento.add_heading(item.get("rotulo") or "Seção", level=1)
            continue
        documento.add_heading(item.get("rotulo") or "Campo", level=2)
        documento.add_paragraph(str(item.get("valor") or "Não informado"))

    nome_profissional = str(
        dados.get("profissional") or "Profissional responsável"
    ).strip()
    espaco_assinaturas = documento.add_paragraph()
    espaco_assinaturas.paragraph_format.space_before = Pt(36)
    tabela_assinaturas = documento.add_table(rows=1, cols=2)
    assinaturas = (
        (nome_profissional, "Assinatura do(a) profissional responsável"),
        (
            str(dados.get("paciente") or "Paciente ou responsável").strip(),
            "Assinatura do paciente ou responsável",
        ),
    )
    for celula, (nome, identificacao) in zip(
        tabela_assinaturas.rows[0].cells, assinaturas
    ):
        linha = celula.paragraphs[0]
        linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        linha.add_run("________________________________")
        nome_assinatura = celula.add_paragraph(nome)
        nome_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nome_assinatura.paragraph_format.keep_with_next = True
        rotulo_assinatura = celula.add_paragraph(identificacao)
        rotulo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for execucao in rotulo_assinatura.runs:
            execucao.font.size = Pt(9)
    documento.save(str(caminho))


def exportar_ficha_pdf(dados: dict, caminho: str) -> None:
    """Gera o PDF usando exatamente os mesmos dados preparados para o Word."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        KeepTogether,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    estilos = getSampleStyleSheet()
    estilo_titulo = ParagraphStyle(
        "FichaTitulo",
        parent=estilos["Title"],
        textColor=colors.HexColor("#075985"),
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    estilo_secao = ParagraphStyle(
        "FichaSecao",
        parent=estilos["Heading2"],
        textColor=colors.HexColor("#075985"),
        backColor=colors.HexColor("#eef8fe"),
        borderPadding=6,
        spaceBefore=12,
        spaceAfter=8,
    )
    estilo_rotulo = ParagraphStyle(
        "FichaRotulo",
        parent=estilos["Heading3"],
        spaceBefore=8,
        spaceAfter=3,
    )
    estilo_resposta = ParagraphStyle(
        "FichaResposta",
        parent=estilos["BodyText"],
        leading=16,
        spaceAfter=6,
    )
    estilo_assinatura = ParagraphStyle(
        "FichaAssinatura",
        parent=estilos["BodyText"],
        alignment=TA_CENTER,
        leading=14,
    )

    elementos = [
        Paragraph(
            html.escape(str(dados.get("modelo") or "Ficha Clínica")),
            estilo_titulo,
        ),
        Paragraph(
            "<b>Paciente:</b> "
            + html.escape(str(dados.get("paciente") or "Não informado")),
            estilos["BodyText"],
        ),
        Paragraph(
            "<b>Data:</b> " + html.escape(str(dados.get("data") or "")),
            estilos["BodyText"],
        ),
    ]
    if dados.get("clinica"):
        elementos.append(Paragraph(
            "<b>Clínica:</b> " + html.escape(str(dados["clinica"])),
            estilos["BodyText"],
        ))
    if dados.get("profissional"):
        elementos.append(Paragraph(
            "<b>Profissional:</b> "
            + html.escape(str(dados["profissional"])),
            estilos["BodyText"],
        ))
    elementos.append(Spacer(1, 8))

    for item in dados.get("itens") or []:
        rotulo = html.escape(str(item.get("rotulo") or "Campo"))
        if item.get("tipo") == "secao":
            elementos.append(Paragraph(rotulo, estilo_secao))
            continue
        valor = html.escape(
            str(item.get("valor") or "Não informado")
        ).replace("\n", "<br/>")
        elementos.append(Paragraph(rotulo, estilo_rotulo))
        elementos.append(Paragraph(valor, estilo_resposta))

    nome_profissional = html.escape(str(
        dados.get("profissional") or "Profissional responsável"
    ).strip())
    nome_paciente = html.escape(str(
        dados.get("paciente") or "Paciente ou responsável"
    ).strip())
    tabela_assinaturas = Table(
        [
            ["", "", ""],
            [
                Paragraph(nome_profissional, estilo_assinatura),
                "",
                Paragraph(nome_paciente, estilo_assinatura),
            ],
            [
                Paragraph(
                    "Assinatura do(a) profissional responsável",
                    estilo_assinatura,
                ),
                "",
                Paragraph(
                    "Assinatura do paciente ou responsável",
                    estilo_assinatura,
                ),
            ],
        ],
        colWidths=[7.4 * cm, 0.8 * cm, 7.4 * cm],
    )
    tabela_assinaturas.setStyle(TableStyle([
        ("LINEABOVE", (0, 0), (0, 0), 0.7, colors.HexColor("#172033")),
        ("LINEABOVE", (2, 0), (2, 0), 0.7, colors.HexColor("#172033")),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elementos.extend([
        Spacer(1, 1.2 * cm),
        KeepTogether([tabela_assinaturas]),
    ])
    documento = SimpleDocTemplate(
        str(caminho),
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title=str(dados.get("modelo") or "Ficha Clínica"),
        author=str(
            dados.get("profissional")
            or dados.get("clinica")
            or "Profissional responsável"
        ),
    )
    documento.build(elementos)


def html_exportacao_ficha(dados: dict) -> str:
    """Entrega o conteúdo seguro usado pelo exportador PDF do Qt."""
    blocos = []
    for item in dados.get("itens") or []:
        rotulo = html.escape(str(item.get("rotulo") or ""))
        if item.get("tipo") == "secao":
            blocos.append(f"<h2>{rotulo}</h2>")
            continue
        valor = html.escape(
            str(item.get("valor") or "Não informado")
        ).replace("\n", "<br>")
        blocos.append(f"<section><h3>{rotulo}</h3><p>{valor}</p></section>")
    profissional = (
        f"<p><strong>Profissional:</strong> "
        f"{html.escape(str(dados.get('profissional')))}</p>"
        if dados.get("profissional") else ""
    )
    clinica = (
        f"<p><strong>Clínica:</strong> "
        f"{html.escape(str(dados.get('clinica')))}</p>"
        if dados.get("clinica") else ""
    )
    return f"""
    <html><head><meta charset="utf-8"><style>
      body {{ font-family: Arial, sans-serif; color: #172033; }}
      header {{ border-bottom: 2px solid #0b8dca; padding-bottom: 12px; }}
      h1 {{ color: #075985; font-size: 24px; }}
      h2 {{ color: #075985; background: #eef8fe; padding: 7px; margin-top: 18px; }}
      h3 {{ font-size: 14px; margin-bottom: 4px; }}
      p {{ font-size: 12px; line-height: 1.45; margin-top: 0; }}
      .assinaturas {{ display: table; margin-top: 48px; width: 100%; }}
      .assinatura {{ display: table-cell; padding: 6px 20px 0;
                     text-align: center; width: 50%; }}
      .linha-assinatura {{ border-top: 1px solid #172033; padding-top: 6px; }}
    </style></head><body>
      <header>
        <h1>{html.escape(str(dados.get('modelo') or 'Ficha Clínica'))}</h1>
        <p><strong>Paciente:</strong> {html.escape(str(dados.get('paciente') or 'Não informado'))}</p>
        <p><strong>Data:</strong> {html.escape(str(dados.get('data') or ''))}</p>
        {clinica}{profissional}
      </header>
      {''.join(blocos) or '<p>Nenhuma resposta preenchida.</p>'}
      <div class="assinaturas">
        <div class="assinatura">
          <p class="linha-assinatura">{html.escape(str(dados.get('profissional') or 'Profissional responsável'))}</p>
          <p>Assinatura do(a) profissional responsável</p>
        </div>
        <div class="assinatura">
          <p class="linha-assinatura">{html.escape(str(dados.get('paciente') or 'Paciente ou responsável'))}</p>
          <p>Assinatura do paciente ou responsável</p>
        </div>
      </div>
    </body></html>
    """
