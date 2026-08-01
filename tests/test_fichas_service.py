from datetime import date

from services.fichas_service import (
    atualizar_respostas_calculadas,
    calcular_idade,
    exportar_ficha_word,
    exportar_ficha_pdf,
    formatar_cpf_ficha,
    formatar_data_ficha,
    formatar_rg_ficha,
    formatar_telefone_ficha,
    html_exportacao_ficha,
    nome_arquivo_exportacao,
    normalizar_estrutura,
    preparar_dados_exportacao,
    respostas_iniciais,
    validar_respostas,
)


def test_normaliza_campos_e_remove_identificador_do_rotulo():
    campos = normalizar_estrutura([
        {
            "tipo": "texto_longo",
            "id": "custom_35f06f7e_prescricoes",
            "label": "35f06f7e prescrições",
            "obrigatorio": True,
        },
        {
            "tipo": "multipla_escolha",
            "label": "Conduta",
            "opcoes": [" Observar ", "", "Encaminhar"],
        },
    ])

    assert campos[0]["label"] == "prescrições"
    assert campos[0]["obrigatorio"] is True
    assert campos[1]["id"].startswith("campo_1_conduta")
    assert campos[1]["opcoes"] == ["Observar", "Encaminhar"]


def test_respostas_iniciais_respeitam_checkbox_e_data():
    campos = normalizar_estrutura([
        {"tipo": "checkbox", "id": "fumante", "label": "Fumante"},
        {
            "tipo": "data",
            "id": "data_consulta",
            "label": "Data",
            "preencher_hoje": True,
        },
    ])

    respostas = respostas_iniciais(campos)

    assert respostas["fumante"] is False
    assert len(respostas["data_consulta"].split("/")) == 3


def test_reconhece_campos_importados_e_relaciona_data_com_idade():
    campos = normalizar_estrutura([
        {"tipo": "data", "id": "data", "label": "Data"},
        {"tipo": "numero", "id": "idade", "label": "Idade"},
        {"tipo": "texto_curto", "id": "celular", "label": "Telefone / celular"},
        {"tipo": "texto_curto", "id": "documento", "label": "CPF"},
    ])

    assert campos[0]["semantica"] == "data_nascimento"
    assert campos[1]["semantica"] == "idade"
    assert campos[1]["calculado_por"] == "data"
    assert campos[1]["somente_leitura"] is True
    assert campos[2]["semantica"] == "telefone"
    assert campos[3]["semantica"] == "cpf"


def test_normaliza_mascaras_de_campos_clinicos():
    assert formatar_data_ficha("16052007") == "16/05/2007"
    assert formatar_telefone_ficha("81991214670") == "(81) 99121-4670"
    assert formatar_cpf_ficha("12345678900") == "123.456.789-00"
    assert formatar_rg_ficha("123456789") == "12.345.678-9"


def test_calcula_idade_e_atualiza_respostas_derivadas():
    campos = normalizar_estrutura([
        {
            "tipo": "data",
            "id": "nascimento",
            "label": "Data de nascimento",
        },
        {"tipo": "numero", "id": "idade", "label": "Idade"},
    ])

    respostas = atualizar_respostas_calculadas(
        campos,
        {"nascimento": "16052007", "idade": ""},
        hoje=date(2026, 7, 31),
    )

    assert respostas["nascimento"] == "16/05/2007"
    assert respostas["idade"] == "19"
    assert calcular_idade("01/08/2007", date(2026, 7, 31)) == "18"
    assert calcular_idade("31/02/2007", date(2026, 7, 31)) == ""


def test_validacao_informa_campos_obrigatorios():
    campos = normalizar_estrutura([
        {
            "tipo": "texto_curto",
            "id": "queixa",
            "label": "Queixa principal",
            "obrigatorio": True,
        },
        {
            "tipo": "texto_longo",
            "id": "observacao",
            "label": "Observação",
        },
    ])

    assert validar_respostas(campos, {"queixa": ""}) == (
        False,
        ["Queixa principal"],
    )
    assert validar_respostas(campos, {"queixa": "Dor lombar"}) == (True, [])


def test_exportacao_preserva_secoes_respostas_e_escapa_html():
    campos = [
        {"tipo": "secao", "label": "Consulta"},
        {"tipo": "texto_longo", "id": "qp", "label": "Queixa"},
        {"tipo": "checkbox", "id": "autorizou", "label": "Autorizou"},
    ]
    dados = preparar_dados_exportacao(
        campos,
        {"qp": "Dor <forte>", "autorizou": True},
        "Maria da Silva",
        "Consulta geral",
        "Clínica Prontu",
        "Dra. Laura",
        "29/07/2026 às 10:00",
    )

    assert dados["itens"][0]["tipo"] == "secao"
    assert dados["itens"][1]["valor"] == "Dor <forte>"
    assert dados["itens"][2]["valor"] == "Sim"
    assert nome_arquivo_exportacao(dados, "pdf").startswith(
        "Ficha_Maria_da_Silva_"
    )
    conteudo = html_exportacao_ficha(dados)
    assert "Dor &lt;forte&gt;" in conteudo
    assert "Clínica Prontu" in conteudo
    assert "Documento gerado pelo Prontu" not in conteudo
    assert "Assinatura do(a) profissional responsável" in conteudo
    assert "Assinatura do paciente ou responsável" in conteudo
    assert "Dra. Laura" in conteudo


def test_exportacao_word_cria_documento_valido(tmp_path):
    from docx import Document

    dados = preparar_dados_exportacao(
        [{"tipo": "texto_curto", "id": "qp", "label": "Queixa"}],
        {"qp": "Avaliação de rotina"},
        "Maria da Silva",
        "Consulta geral",
        "Clínica Horizonte",
        "Dra. Laura",
    )
    destino = tmp_path / "ficha.docx"

    exportar_ficha_word(dados, destino)

    assert destino.read_bytes().startswith(b"PK")
    assert destino.stat().st_size > 1000
    documento = Document(destino)
    texto = "\n".join([
        *(paragrafo.text for paragrafo in documento.paragraphs),
        *(
            paragrafo.text
            for tabela in documento.tables
            for linha in tabela.rows
            for celula in linha.cells
            for paragrafo in celula.paragraphs
        ),
    ])
    assert "Queixa" in texto
    assert "Avaliação de rotina" in texto
    assert "Documento gerado pelo Prontu" not in texto
    assert "Assinatura do(a) profissional responsável" in texto
    assert "Assinatura do paciente ou responsável" in texto
    assert "Dra. Laura" in texto
    assert "Maria da Silva" in texto
    assert len(documento.inline_shapes) == 0


def test_exportacao_pdf_inclui_rotulos_e_respostas(tmp_path):
    from pypdf import PdfReader

    dados = preparar_dados_exportacao(
        [{"tipo": "texto_longo", "id": "qp", "label": "Queixa principal"}],
        {"qp": "Dor lombar há três dias"},
        "Maria da Silva",
        "Consulta geral",
        "Clínica Horizonte",
        "Dra. Laura",
    )
    destino = tmp_path / "ficha.pdf"

    exportar_ficha_pdf(dados, destino)

    texto = "\n".join(
        pagina.extract_text() or "" for pagina in PdfReader(destino).pages
    )
    assert "Queixa principal" in texto
    assert "Dor lombar" in texto
    assert "Documento gerado pelo Prontu" not in texto
    assert "Assinatura do(a) profissional responsável" in texto
    assert "Assinatura do paciente ou responsável" in texto
    assert "Dra. Laura" in texto
    assert PdfReader(destino).metadata.author == "Dra. Laura"
