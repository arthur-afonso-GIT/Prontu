"""Testes da engine de parsing: extração e classificação de PDF/DOCX."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from forms.field_schema import FieldType
from parsers.parser_engine import ParserEngine, ParserEngineError


@pytest.fixture
def docx_ficha_completa(tmp_path: Path) -> Path:
    """Cria um DOCX sintético com título, campos simples, textarea, tabela e score."""
    doc = Document()
    doc.add_paragraph("FICHA DE AVALIAÇÃO NUTRICIONAL")
    doc.add_paragraph("Nome:")
    doc.add_paragraph("Telefone:")
    doc.add_paragraph("Queixa Principal:")

    tabela = doc.add_table(rows=2, cols=3)
    tabela.rows[0].cells[0].text = "Data"
    tabela.rows[0].cells[1].text = "Peso (kg)"
    tabela.rows[0].cells[2].text = "IMC"
    tabela.rows[1].cells[0].text = "01/01/2026"
    tabela.rows[1].cells[1].text = "78.5"
    tabela.rows[1].cells[2].text = "24.3"

    doc.add_paragraph("ESCORE CHA2DS2-VASc")
    score_table = doc.add_table(rows=2, cols=2)
    score_table.rows[0].cells[0].text = "Critério"
    score_table.rows[0].cells[1].text = "Pontos"
    score_table.rows[1].cells[0].text = "Hipertensão arterial"
    score_table.rows[1].cells[1].text = "1"

    caminho = tmp_path / "ficha_teste.docx"
    doc.save(str(caminho))
    return caminho


class TestParserEngine:
    def test_formato_nao_suportado_levanta_erro(self, tmp_path: Path):
        arquivo_invalido = tmp_path / "arquivo.txt"
        arquivo_invalido.write_text("conteudo")

        engine = ParserEngine()
        with pytest.raises(ParserEngineError):
            engine.parse_document(arquivo_invalido)

    def test_detecta_campo_nome_e_telefone(self, docx_ficha_completa: Path):
        engine = ParserEngine()
        campos = engine.parse_document(docx_ficha_completa)

        ids_detectados = {campo.id for campo in campos}
        assert "nome" in ids_detectados
        assert "telefone" in ids_detectados

    def test_campo_queixa_principal_e_textarea(self, docx_ficha_completa: Path):
        engine = ParserEngine()
        campos = engine.parse_document(docx_ficha_completa)

        campo_qp = next(c for c in campos if c.id == "qp")
        assert campo_qp.tipo == FieldType.TEXTAREA

    def test_titulo_e_detectado(self, docx_ficha_completa: Path):
        engine = ParserEngine()
        campos = engine.parse_document(docx_ficha_completa)

        titulos = [c for c in campos if c.tipo in (FieldType.TITULO, FieldType.SUBTITULO)]
        assert len(titulos) >= 1

    def test_tabela_antropometrica_detectada_com_colunas_numericas(self, docx_ficha_completa: Path):
        engine = ParserEngine()
        campos = engine.parse_document(docx_ficha_completa)

        tabelas = [c for c in campos if c.tipo == FieldType.TABELA]
        assert len(tabelas) == 1

        coluna_peso = next(col for col in tabelas[0].colunas if "peso" in col.label.lower())
        assert coluna_peso.tipo == FieldType.NUMERO
        assert coluna_peso.unidade == "kg"

    def test_score_cha2ds2vasc_reconhecido_pelo_nome(self, docx_ficha_completa: Path):
        engine = ParserEngine()
        campos = engine.parse_document(docx_ficha_completa)

        scores = [c for c in campos if c.tipo == FieldType.SCORE]
        assert len(scores) == 1
        assert "CHA2DS2-VASc" in scores[0].label
        # Score curado deve ter os 8 critérios completos da escala, não apenas
        # o único item presente na tabela de exemplo do documento de teste.
        assert len(scores[0].itens_score) == 8


class TestScoreCalculation:
    def test_calculo_de_pontuacao_total(self):
        from parsers.score_detector import KNOWN_SCORES, ScoreDetector

        detector = ScoreDetector()
        definicao = KNOWN_SCORES["cha2ds2vasc"]
        campo_score = detector.build_score_field(definicao, ordem=0)

        itens_marcados = {"hipertensao", "idade_75"}  # 1 + 2 = 3 pontos
        total = ScoreDetector.calculate_score_total(campo_score.itens_score, itens_marcados)
        assert total == 3

    def test_classificacao_de_risco_por_faixa(self):
        from parsers.score_detector import KNOWN_SCORES, ScoreDetector

        detector = ScoreDetector()
        definicao = KNOWN_SCORES["cha2ds2vasc"]
        campo_score = detector.build_score_field(definicao, ordem=0)

        faixa = ScoreDetector.classify_risk(campo_score.faixas_risco, 0)
        assert faixa.classificacao == "Risco baixo"

        faixa_alta = ScoreDetector.classify_risk(campo_score.faixas_risco, 4)
        assert faixa_alta.classificacao == "Risco alto"
