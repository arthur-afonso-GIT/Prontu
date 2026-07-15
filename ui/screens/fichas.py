import os
import json
import uuid
import mimetypes
import re
import unicodedata
import html
from datetime import datetime
from PySide6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel, 
                               QLineEdit, QPushButton, QComboBox, QScrollArea, 
                               QFrame, QCheckBox, QTextEdit, QFileDialog, QMessageBox, QListView, QDialog,
                               QDateEdit, QRadioButton, QButtonGroup, QCompleter)
from PySide6.QtGui import QPixmap, QDesktopServices, QColor, QDoubleValidator, QTextDocument
from PySide6.QtCore import Qt, QUrl, QDate
from PySide6.QtPrintSupport import QPrinter

try:
    from docx import Document
    DOCX_DISPONIVEL = True
except ImportError:
    Document = None
    DOCX_DISPONIVEL = False

try:
    from pypdf import PdfReader
    PDF_DISPONIVEL = True
except ImportError:
    PdfReader = None
    PDF_DISPONIVEL = False

# Extensões aceitas para anexo em uma ficha preenchida
EXTENSOES_ANEXO_ACEITAS = "Arquivos Suportados (*.jpg *.jpeg *.png *.webp *.pdf)"
NOME_BUCKET_ANEXOS = "fichas-anexos"


class CustomInputDialog(QDialog):
    def __init__(self, titulo, mensagem, parent=None):
        super().__init__(parent)
        self.setWindowTitle(titulo)
        self.setFixedWidth(420)
        self.setStyleSheet("background-color: #ffffff;")
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)
        
        self.label = QLabel(mensagem)
        self.label.setWordWrap(True)
        self.label.setStyleSheet("color: #0f172a !important; font-size: 13px !important; font-weight: 500 !important; background: transparent;")
        layout.addWidget(self.label)
        
        self.input_field = QLineEdit()
        self.input_field.setStyleSheet("""
            QLineEdit { 
                background-color: #ffffff !important; 
                color: #0f172a !important; 
                border: 1px solid #cbd5e1 !important; 
                border-radius: 6px; 
                padding: 8px; 
                font-size: 13px; 
            }
            QLineEdit:focus { border: 1px solid #0284c7 !important; }
        """)
        layout.addWidget(self.input_field)
        
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)
        
        self.btn_cancelar = QPushButton("Cancelar")
        self.btn_cancelar.setStyleSheet("QPushButton { background-color: #f1f5f9; color: #0f172a; border: 1px solid #cbd5e1; padding: 7px 14px; border-radius: 5px; font-weight: bold; } QPushButton:hover { background-color: #e2e8f0; }")
        self.btn_cancelar.clicked.connect(self.reject)
        
        self.btn_confirmar = QPushButton("Confirmar")
        self.btn_confirmar.setStyleSheet("QPushButton { background-color: #0284c7; color: white; padding: 7px 14px; border-radius: 5px; font-weight: bold; border: none; } QPushButton:hover { background-color: #0369a1; }")
        self.btn_confirmar.clicked.connect(self.accept)
        
        btn_layout.addStretch()
        btn_layout.addWidget(self.btn_cancelar)
        btn_layout.addWidget(self.btn_confirmar)
        layout.addLayout(btn_layout)

    def get_text(self):
        return self.input_field.text().strip()


class FichasScreen(QWidget):
    def __init__(self, database_instancia):
        super().__init__()
        
        self.db = database_instancia
        self._formulario_sujo = False
        self.modelo_atual_campos = [] 
        self.widgets_dinamicos = {}   
        self.modo_criacao = False 
        self._nome_modelo_importado = None
        self.ficha_em_edicao_id = None
        self._data_atendimento_original = None
        self._anexos_existentes = []
        self.arquivos_anexados = []  # Lista de caminhos locais pendentes de upload (limpa após salvar)
        
        self.main_layout = QHBoxLayout(self)
        self.main_layout.setContentsMargins(20, 20, 20, 20)
        self.main_layout.setSpacing(20)
        
        # --- COLUNA ESQUERDA: Configurações ---
        self.left_panel = QFrame()
        self.left_panel.setFixedWidth(320)
        self.left_panel.setStyleSheet("""
            QFrame { 
                background-color: white; 
                border: 1px solid #e2e8f0; 
                border-radius: 12px; 
            } 
            QLabel { 
                color: #0f172a !important; 
                font-weight: 500; 
                font-size: 13px; 
                border: none; 
                background-color: transparent;
            }
            QComboBox { 
                background-color: #ffffff !important; 
                color: #0f172a !important; 
                border: 1px solid #cbd5e1 !important; 
                border-radius: 6px; 
                padding: 6px; 
            }
        """)
        self.left_layout = QVBoxLayout(self.left_panel)
        self.left_layout.setSpacing(12)
        self.left_layout.setContentsMargins(20, 20, 20, 20)
        
        titulo_painel = QLabel("📋 Prontuário & Anamnese")
        titulo_painel.setStyleSheet("font-size: 16px; font-weight: bold; color: #0f172a !important;")
        self.left_layout.addWidget(titulo_painel)
        
        self.left_layout.addWidget(QLabel("1. Selecione o Paciente:"))
        self.combo_paciente = QComboBox()
        self.combo_paciente.setView(QListView())
        self.combo_paciente.view().setStyleSheet("QListView { background-color: #ffffff !important; color: #0f172a !important; selection-background-color: #0284c7; }")
        self.combo_paciente.setEditable(True)
        self.combo_paciente.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.combo_paciente.lineEdit().setPlaceholderText("Digite o nome do paciente...")
        self.combo_paciente.lineEdit().setClearButtonEnabled(True)
        self.busca_paciente = QCompleter(self.combo_paciente.model(), self.combo_paciente)
        self.busca_paciente.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        self.busca_paciente.setFilterMode(Qt.MatchFlag.MatchContains)
        self.busca_paciente.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
        self.combo_paciente.setCompleter(self.busca_paciente)
        self.left_layout.addWidget(self.combo_paciente)
        
        self.left_layout.addWidget(QLabel("2. Modelo de Ficha Clínica:"))
        self.combo_modelo = QComboBox()
        self.combo_modelo.setView(QListView())
        self.combo_modelo.view().setStyleSheet("QListView { background-color: #ffffff !important; color: #0f172a !important; selection-background-color: #0284c7; }")
        self.left_layout.addWidget(self.combo_modelo)
        
        botoes_importar_layout = QHBoxLayout()
        botoes_importar_layout.setSpacing(8)

        self.btn_importar = QPushButton("📥 Word (.docx)")
        self.btn_importar.setStyleSheet("""
            QPushButton { background-color: #f1f5f9; color: #0f172a; border: 1px solid #cbd5e1; padding: 10px; font-weight: bold; border-radius: 6px; }
            QPushButton:hover { background-color: #e2e8f0; }
        """)
        self.btn_importar.clicked.connect(self.importar_modelo_word)
        botoes_importar_layout.addWidget(self.btn_importar)

        self.btn_importar_pdf = QPushButton("📥 PDF")
        self.btn_importar_pdf.setStyleSheet("""
            QPushButton { background-color: #f1f5f9; color: #0f172a; border: 1px solid #cbd5e1; padding: 10px; font-weight: bold; border-radius: 6px; }
            QPushButton:hover { background-color: #e2e8f0; }
        """)
        self.btn_importar_pdf.clicked.connect(self.importar_modelo_pdf)
        botoes_importar_layout.addWidget(self.btn_importar_pdf)

        self.left_layout.addLayout(botoes_importar_layout)
        
        self.btn_criar_modelo = QPushButton("🛠️ Montar Novo Modelo")
        self.btn_criar_modelo.setStyleSheet("""
            QPushButton { background-color: #f0fdf4; color: #166534; border: 1px solid #bbf7d0; padding: 10px; font-weight: bold; border-radius: 6px; }
            QPushButton:hover { background-color: #dcfce7; }
        """)
        self.btn_criar_modelo.clicked.connect(self.iniciar_criacao_modelo)
        self.left_layout.addWidget(self.btn_criar_modelo)

        acoes_modelo = QHBoxLayout()
        self.btn_editar_modelo = QPushButton("Editar modelo")
        self.btn_editar_modelo.setStyleSheet(
            "QPushButton { background-color: #eff6ff; color: #1d4ed8; border: 1px solid #93c5fd; padding: 8px; font-weight: bold; border-radius: 6px; }"
        )
        self.btn_editar_modelo.clicked.connect(self.editar_modelo_selecionado)
        self.btn_excluir_modelo = QPushButton("Excluir modelo")
        self.btn_excluir_modelo.setStyleSheet(
            "QPushButton { background-color: #fff7ed; color: #c2410c; border: 1px solid #fdba74; padding: 8px; font-weight: bold; border-radius: 6px; }"
        )
        self.btn_excluir_modelo.clicked.connect(self.excluir_modelo_selecionado)
        acoes_modelo.addWidget(self.btn_editar_modelo)
        acoes_modelo.addWidget(self.btn_excluir_modelo)
        self.left_layout.addLayout(acoes_modelo)

        # --- SEÇÃO DE ANEXOS (fotos/PDFs vinculados ao atendimento atual) ---
        sep_anexos = QFrame()
        sep_anexos.setStyleSheet("background-color: #e2e8f0; max-height: 1px; border: none; margin: 8px 0;")
        self.left_layout.addWidget(sep_anexos)

        self.left_layout.addWidget(QLabel("3. Anexos do Atendimento:"))

        self.anexos_scroll = QScrollArea()
        self.anexos_scroll.setWidgetResizable(True)
        self.anexos_scroll.setFixedHeight(90)
        self.anexos_scroll.setStyleSheet("QScrollArea { border: 1px solid #e2e8f0; border-radius: 6px; background-color: #f8fafc; }")
        self.anexos_strip_widget = QWidget()
        self.anexos_strip_widget.setStyleSheet("background-color: #f8fafc;")
        self.anexos_strip_layout = QHBoxLayout(self.anexos_strip_widget)
        self.anexos_strip_layout.setContentsMargins(6, 6, 6, 6)
        self.anexos_strip_layout.setSpacing(6)
        self.anexos_strip_layout.addStretch()
        self.anexos_scroll.setWidget(self.anexos_strip_widget)
        self.left_layout.addWidget(self.anexos_scroll)

        self.btn_anexar_arquivo = QPushButton("📎 Anexar Foto ou PDF")
        self.btn_anexar_arquivo.setStyleSheet("""
            QPushButton { background-color: #fffbeb; color: #92400e; border: 1px solid #fde68a; padding: 8px; font-weight: bold; border-radius: 6px; font-size: 12px; }
            QPushButton:hover { background-color: #fef3c7; }
        """)
        self.btn_anexar_arquivo.clicked.connect(self.anexar_arquivos)
        self.left_layout.addWidget(self.btn_anexar_arquivo)
        
        self.left_layout.addStretch()
        
        self.btn_salvar_atendimento = QPushButton("💾 Salvar Ficha Preenchida")
        self.btn_salvar_atendimento.setStyleSheet("""
            QPushButton { background-color: #0284c7; color: white; padding: 12px; font-weight: bold; border-radius: 6px; font-size: 14px; border: none; }
            QPushButton:hover { background-color: #0369a1; }
        """)
        self.btn_salvar_atendimento.clicked.connect(self.salvar_ficha_preenchida)
        self.left_layout.addWidget(self.btn_salvar_atendimento)

        exportacoes_layout = QHBoxLayout()
        exportacoes_layout.setSpacing(6)
        self.btn_exportar_word = QPushButton("Exportar Word")
        self.btn_exportar_pdf = QPushButton("Exportar PDF")
        for botao in (self.btn_exportar_word, self.btn_exportar_pdf):
            botao.setStyleSheet(
                "QPushButton { background: #f1f5f9; color: #1e3a8a; border: 1px solid #bfdbfe; "
                "border-radius: 6px; padding: 7px; font-weight: bold; font-size: 12px; } "
                "QPushButton:hover { background: #dbeafe; }"
            )
        self.btn_exportar_word.clicked.connect(self.exportar_ficha_word)
        self.btn_exportar_pdf.clicked.connect(self.exportar_ficha_pdf)
        exportacoes_layout.addWidget(self.btn_exportar_word)
        exportacoes_layout.addWidget(self.btn_exportar_pdf)
        self.left_layout.addLayout(exportacoes_layout)

        self.btn_cancelar_edicao = QPushButton("Cancelar edição da ficha")
        self.btn_cancelar_edicao.setVisible(False)
        self.btn_cancelar_edicao.setStyleSheet(
            "QPushButton { background-color: #f1f5f9; color: #334155; border: 1px solid #cbd5e1; padding: 7px; font-weight: bold; border-radius: 6px; }"
        )
        self.btn_cancelar_edicao.clicked.connect(self.cancelar_edicao_ficha)
        self.left_layout.addWidget(self.btn_cancelar_edicao)

        self.lbl_status_operacao = QLabel("")
        self.lbl_status_operacao.setWordWrap(True)
        self.lbl_status_operacao.setStyleSheet("color: #15803d; font-size: 12px; padding-top: 4px;")
        self.left_layout.addWidget(self.lbl_status_operacao)
        
        self.main_layout.addWidget(self.left_panel)
        
        # --- COLUNA DIREITA: Container Dinâmico ---
        self.right_container = QWidget()
        self.right_layout = QVBoxLayout(self.right_container)
        self.right_layout.setContentsMargins(0, 0, 0, 0)
        
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setStyleSheet("QScrollArea { border: 1px solid #e2e8f0; border-radius: 12px; background-color: white; }")
        
        self.scroll_content = QWidget()
        self.scroll_content.setObjectName("ScrollContent")
        self.scroll_content.setStyleSheet("#ScrollContent { background-color: white; }")
        
        self.dinamic_form_layout = QVBoxLayout(self.scroll_content)
        self.dinamic_form_layout.setSpacing(14)
        self.dinamic_form_layout.setContentsMargins(25, 25, 25, 25)
        
        self.scroll_area.setWidget(self.scroll_content)
        self.right_layout.addWidget(self.scroll_area)
        
        self.main_layout.addWidget(self.right_container, stretch=3)
        
        # Inicialização Segura baseada no Supabase
        self.carregar_modelos_iniciais_combo()
        self.carregar_pacientes_combo()
        self.gerar_modelo_padrao()
        self.renderizar_anexos_thumbnails()
        
        self.combo_modelo.currentTextChanged.connect(self.alterar_modelo_ficha)

    def carregar_pacientes_combo(self):
        self.combo_paciente.clear()
        if not self.db.supabase:
            return
        try:
            # Seleciona os pacientes vinculados ao consultório logado
            resposta = self.db.supabase.table("pacientes")\
                .select("id, nome")\
                .eq("consultorio_id", self.db.consultorio_id)\
                .is_("deleted_at", "null")\
                .order("nome", desc=False)\
                .execute()
                
            if resposta.data:
                for row in resposta.data:
                    self.combo_paciente.addItem(f"👤 {row['nome']} (ID: {row['id']})", row['id'])
            else:
                self.combo_paciente.addItem("Nenhum paciente cadastrado")
        except Exception as e:
            print(f"Erro ao carregar pacientes do Supabase: {e}")
            self.combo_paciente.addItem("Nenhum paciente cadastrado")

    def carregar_modelos_iniciais_combo(self):
        """Busca modelos de ficha cadastrados na nuvem para o consultório logado."""
        self.combo_modelo.clear()
        self.combo_modelo.addItem("Ficha de Consulta Geral (Padrão)")
        if not self.db.supabase:
            return
        try:
            resposta = self.db.supabase.table("modelos_fichas")\
                .select("nome_modelo")\
                .eq("consultorio_id", self.db.consultorio_id)\
                .order("id", desc=True)\
                .execute()
                
            if resposta.data:
                for m in resposta.data:
                    self.combo_modelo.addItem(m["nome_modelo"])
        except Exception as e:
            print(f"Erro ao carregar modelos do Supabase: {e}")

    def alterar_modelo_ficha(self, nome_modelo):
        if self.modo_criacao or not nome_modelo:
            return
        if "Padrão" in nome_modelo:
            self.gerar_modelo_padrao()
        else:
            if not self.db.supabase:
                return
            try:
                resposta = self.db.supabase.table("modelos_fichas")\
                    .select("estrutura_json")\
                    .eq("consultorio_id", self.db.consultorio_id)\
                    .eq("nome_modelo", nome_modelo)\
                    .maybe_single()\
                    .execute()
                    
                if resposta.data:
                    self.modelo_atual_campos = json.loads(resposta.data["estrutura_json"])
                    self.renderizar_formulario_dinamico()
            except Exception as e:
                print(f"Erro ao obter modelo de ficha: {e}")

    def editar_modelo_selecionado(self):
        nome_modelo = self.combo_modelo.currentText()
        if not nome_modelo or "Padrão" in nome_modelo:
            self.exibir_popup("aviso", "Modelo padrão", "Crie uma cópia para editar o modelo padrão.")
            return
        try:
            resposta = self.db.supabase.table("modelos_fichas")\
                .select("estrutura_json")\
                .eq("consultorio_id", self.db.consultorio_id)\
                .eq("nome_modelo", nome_modelo)\
                .maybe_single().execute()
            if resposta.data:
                self.iniciar_criacao_modelo(json.loads(resposta.data["estrutura_json"]), origem=nome_modelo)
        except Exception as e:
            self.exibir_popup("erro", "Não foi possível abrir", "Não foi possível abrir esse modelo para edição.")

    def excluir_modelo_selecionado(self):
        nome_modelo = self.combo_modelo.currentText()
        if not nome_modelo or "Padrão" in nome_modelo:
            self.exibir_popup("aviso", "Modelo padrão", "O modelo padrão não pode ser excluído.")
            return
        confirmar = QMessageBox.question(
            self, "Excluir modelo", f"Excluir o modelo '{nome_modelo}'? Fichas já preenchidas não serão apagadas.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if confirmar != QMessageBox.StandardButton.Yes:
            return
        try:
            self.db.supabase.table("modelos_fichas").delete()\
                .eq("consultorio_id", self.db.consultorio_id)\
                .eq("nome_modelo", nome_modelo).execute()
            self.carregar_modelos_iniciais_combo()
            self.gerar_modelo_padrao()
        except Exception:
            self.exibir_popup("erro", "Não foi possível excluir", "O modelo não foi excluído.")

    def exibir_popup(self, tipo, titulo, mensagem):
        if tipo == "info":
            self.lbl_status_operacao.setText(mensagem)
            return
        msg = QMessageBox(self)
        if tipo == "info": msg.setIcon(QMessageBox.Information)
        elif tipo == "aviso": msg.setIcon(QMessageBox.Warning)
        elif tipo == "erro": msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(titulo)
        msg.setText(mensagem)
        msg.setStyleSheet("""
            QMessageBox { background-color: #ffffff !important; }
            QLabel { color: #0f172a !important; font-size: 13px !important; font-weight: 500 !important; }
            QPushButton { background-color: #0284c7 !important; color: white !important; border-radius: 4px; padding: 6px 14px; min-width: 70px; }
            QPushButton:hover { background-color: #0369a1 !important; }
        """)
        msg.exec()

    def limpar_layout_completamente(self, layout):
        if layout is not None:
            while layout.count():
                item = layout.takeAt(0)
                widget = item.widget()
                if widget is not None:
                    widget.setParent(None)
                    widget.deleteLater()
                else:
                    sub_layout = item.layout()
                    if sub_layout is not None:
                        self.limpar_layout_completamente(sub_layout)

    def gerar_modelo_padrao(self):
        self.modo_criacao = False
        self.modelo_atual_campos = [
            {"tipo": "secao", "label": "HISTÓRICO DA CONSULTA ATUAL"},
            {"tipo": "texto_longo", "label": "Queixa Principal (QP)", "id": "qp"},
            {"tipo": "texto_longo", "label": "Histórico da Doença Atual (HDA)", "id": "hda"},
            {"tipo": "secao", "label": "EXAME FÍSICO & SINAIS VITAIS"},
            {"tipo": "texto_curto", "label": "Pressão Arterial (PA)", "id": "pa", "placeholder": "120x80 mmHg"},
            {"tipo": "texto_curto", "label": "Frequência Cardíaca (FC)", "id": "fc", "placeholder": "75 bpm"},
            {"tipo": "secao", "label": "CONDUTA MÉDICA"},
            {"tipo": "texto_longo", "label": "Prescrição / Orientações Passadas", "id": "prescricao"}
        ]
        self.renderizar_formulario_dinamico()

    def iniciar_criacao_modelo(self, campos_iniciais=None, origem="Novo modelo"):
        self.modo_criacao = True
        self._nome_modelo_importado = origem
        self.modelo_atual_campos = list(campos_iniciais or [])
        self.widgets_dinamicos.clear()
        
        self.limpar_layout_completamente(self.dinamic_form_layout)
            
        lbl_info = QLabel("🛠️ Construtor de Ficha Personalizada (Modo Preview)")
        lbl_info.setStyleSheet("font-size: 18px; font-weight: bold; color: #0284c7; margin-bottom: 2px;")
        self.dinamic_form_layout.addWidget(lbl_info)
        
        lbl_sub = QLabel("Clique nos botões para adicionar campos. O formulário abaixo atualiza em tempo real:")
        lbl_sub.setStyleSheet("color: #64748b; font-size: 13px; margin-bottom: 12px;")
        self.dinamic_form_layout.addWidget(lbl_sub)
        
        btn_layout = QVBoxLayout()
        btn_layout.setSpacing(8)
        primeira_linha_botoes = QHBoxLayout()
        segunda_linha_botoes = QHBoxLayout()
        
        btn_add_secao = QPushButton("+ Seção (Título)")
        btn_add_curto = QPushButton("+ Texto Curto")
        btn_add_longo = QPushButton("+ Texto Longo")
        btn_add_check = QPushButton("+ Caixa de Seleção")
        btn_add_numero = QPushButton("+ Número")
        btn_add_data = QPushButton("+ Data")
        btn_add_multipla = QPushButton("+ Múltipla Escolha")
        btn_modelo_inicial = QPushButton("Usar modelo inicial")
        
        estilo_botoes_add = """
            QPushButton { background-color: #f8fafc; color: #334155; border: 1px solid #cbd5e1; border-radius: 6px; padding: 10px; font-weight: bold; font-size: 12px;}
            QPushButton:hover { background-color: #f1f5f9; border: 1px solid #94a3b8; }
        """
        botoes_add = [btn_add_secao, btn_add_curto, btn_add_longo, btn_add_check, btn_add_numero, btn_add_data, btn_add_multipla, btn_modelo_inicial]
        for indice_botao, b in enumerate(botoes_add):
            b.setStyleSheet(estilo_botoes_add)
            if indice_botao < 4:
                primeira_linha_botoes.addWidget(b)
            else:
                segunda_linha_botoes.addWidget(b)
        btn_layout.addLayout(primeira_linha_botoes)
        btn_layout.addLayout(segunda_linha_botoes)
            
        btn_add_secao.clicked.connect(lambda: self.adicionar_elemento_rascunho("secao"))
        btn_add_curto.clicked.connect(lambda: self.adicionar_elemento_rascunho("texto_curto"))
        btn_add_longo.clicked.connect(lambda: self.adicionar_elemento_rascunho("texto_longo"))
        btn_add_check.clicked.connect(lambda: self.adicionar_elemento_rascunho("checkbox"))
        btn_add_numero.clicked.connect(lambda: self.adicionar_elemento_rascunho("numero"))
        btn_add_data.clicked.connect(lambda: self.adicionar_elemento_rascunho("data"))
        btn_add_multipla.clicked.connect(lambda: self.adicionar_elemento_rascunho("multipla_escolha"))
        btn_modelo_inicial.clicked.connect(self.aplicar_modelo_inicial_rascunho)
        
        self.dinamic_form_layout.addLayout(btn_layout)
        
        sep = QFrame()
        sep.setStyleSheet("background-color: #cbd5e1; max-height: 2px; border: none; margin: 15px 0 5px 0;")
        self.dinamic_form_layout.addWidget(sep)
        
        self.preview_layout = QVBoxLayout()
        self.preview_layout.setSpacing(12)
        self.dinamic_form_layout.addLayout(self.preview_layout)
        
        sep_bottom = QFrame()
        sep_bottom.setStyleSheet("background-color: #e2e8f0; max-height: 1px; border: none; margin-top: 20px;")
        self.dinamic_form_layout.addWidget(sep_bottom)
        
        acoes_layout = QHBoxLayout()
        acoes_layout.setContentsMargins(0, 10, 0, 0)
        
        btn_cancelar = QPushButton("Cancelar")
        btn_cancelar.setStyleSheet("QPushButton { background-color: #ef4444; color: white; padding: 11px 20px; border-radius: 6px; font-weight: bold; border: none; } QPushButton:hover { background-color: #dc2626; }")
        btn_cancelar.clicked.connect(self.gerar_modelo_padrao)
        
        btn_finalizar = QPushButton("💾 Concluir e Salvar Modelo")
        btn_finalizar.setStyleSheet("QPushButton { background-color: #16a34a; color: white; padding: 11px 24px; border-radius: 6px; font-weight: bold; border: none; } QPushButton:hover { background-color: #15803d; }")
        btn_finalizar.clicked.connect(self.salvar_modelo_customizado_db)
        
        acoes_layout.addWidget(btn_cancelar)
        acoes_layout.addStretch()
        acoes_layout.addWidget(btn_finalizar)
        self.dinamic_form_layout.addLayout(acoes_layout)
        
        self.atualizar_visualizacao_preview()

    def aplicar_modelo_inicial_rascunho(self):
        """Oferece uma base editável, evitando começar um modelo do zero."""
        if self.modelo_atual_campos:
            resposta = QMessageBox.question(
                self, "Substituir rascunho", "Trocar os campos atuais pelo modelo inicial?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if resposta != QMessageBox.StandardButton.Yes:
                return
        self.modelo_atual_campos = [
            {"tipo": "secao", "label": "Identificacao"},
            {"tipo": "texto_curto", "label": "Motivo da consulta", "id": "motivo_consulta"},
            {"tipo": "secao", "label": "Avaliacao"},
            {"tipo": "texto_longo", "label": "Historico e observacoes", "id": "historico_observacoes"},
            {"tipo": "numero", "label": "Peso", "id": "peso", "unidade": "kg"},
            {"tipo": "data", "label": "Data do atendimento", "id": "data_atendimento"},
            {"tipo": "secao", "label": "Plano"},
            {"tipo": "texto_longo", "label": "Conduta e orientacoes", "id": "conduta_orientacoes"},
        ]
        self.atualizar_visualizacao_preview()

    def adicionar_elemento_rascunho(self, tipo):
        titulos = {
            "secao": "Nova Seção",
            "numero": "Novo Campo Numérico",
            "data": "Novo Campo de Data",
            "multipla_escolha": "Nova Pergunta de Múltipla Escolha",
        }
        mensagens = {
            "secao": "Digite o título da seção divisor:",
            "numero": "Digite o nome do campo (ex: Peso, Altura, Idade):",
            "data": "Digite o nome do campo de data (ex: Data do Exame):",
            "multipla_escolha": "Digite o texto da pergunta:",
        }
        dialog_tit = titulos.get(tipo, "Novo Campo")
        dialog_msg = mensagens.get(tipo, "Digite o nome da pergunta/campo (ex: Histórico Familiar):")

        dial = CustomInputDialog(dialog_tit, dialog_msg, self)
        if dial.exec() == QDialog.Accepted:
            texto = dial.get_text()
            if not texto: return
        else:
            return
            
        texto_limpo = "".join(c for c in texto if c.isalnum()).lower()
        id_campo = f"custom_{int(datetime.now().timestamp())}_{texto_limpo}"
        novo_campo = {"tipo": tipo, "label": texto, "id": id_campo}

        if tipo == "numero":
            dial_unidade = CustomInputDialog("Unidade (opcional)", "Ex: kg, cm, anos, mmHg — deixe em branco se não tiver:", self)
            if dial_unidade.exec() == QDialog.Accepted:
                unidade = dial_unidade.get_text().strip()
                if unidade:
                    novo_campo["unidade"] = unidade

        elif tipo == "multipla_escolha":
            dial_opcoes = CustomInputDialog("Opções de Resposta", "Digite as opções separadas por vírgula (ex: Sim, Não, Não sei):", self)
            if dial_opcoes.exec() != QDialog.Accepted:
                return
            texto_opcoes = dial_opcoes.get_text().strip()
            opcoes = [o.strip() for o in texto_opcoes.split(",") if o.strip()]
            if len(opcoes) < 2:
                self.exibir_popup("aviso", "Opções insuficientes", "Informe pelo menos 2 opções separadas por vírgula.")
                return
            novo_campo["opcoes"] = opcoes

        self.modelo_atual_campos.append(novo_campo)
        self.atualizar_visualizacao_preview()

    def atualizar_visualizacao_preview(self):
        self.limpar_layout_completamente(self.preview_layout)
        
        if not self.modelo_atual_campos:
            lbl_vazio = QLabel("(Nenhum campo adicionado ainda. Monte sua estrutura clicando nos botões acima...)")
            lbl_vazio.setStyleSheet("color: #94a3b8; font-style: italic; font-size: 13px; padding: 20px; text-align: center;")
            lbl_vazio.setAlignment(Qt.AlignCenter)
            self.preview_layout.addWidget(lbl_vazio)
            return

        estilo_label = "color: #334155 !important; font-weight: bold; font-size: 13px; margin-top: 4px; background-color: transparent;"
        estilo_input_curto = "QLineEdit { background-color: #f8fafc !important; color: #0f172a !important; border: 1px solid #cbd5e1 !important; border-radius: 6px; padding: 8px 12px; font-size: 13px; }"
        estilo_input_longo = "QTextEdit { background-color: #f8fafc !important; color: #0f172a !important; border: 1px solid #cbd5e1 !important; border-radius: 6px; padding: 8px 12px; font-size: 13px; }"
        estilo_btn_ferramenta = "QPushButton { background-color: #ffffff; color: #475569; border: 1px solid #cbd5e1; border-radius: 4px; font-size: 10px; padding: 2px; } QPushButton:hover { background-color: #f1f5f9; }"

        total_campos = len(self.modelo_atual_campos)
        cabecalho_preview = QFrame()
        cabecalho_preview.setStyleSheet(
            "QFrame { background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 6px; }"
        )
        cabecalho_layout = QHBoxLayout(cabecalho_preview)
        cabecalho_layout.setContentsMargins(12, 8, 12, 8)
        titulo_preview = QLabel("Pre-visualizacao da ficha")
        titulo_preview.setStyleSheet("font-weight: bold; color: #1e3a8a; border: none;")
        contador_preview = QLabel(f"{total_campos} elemento(s)")
        contador_preview.setStyleSheet("color: #475569; border: none;")
        cabecalho_layout.addWidget(titulo_preview)
        cabecalho_layout.addStretch()
        cabecalho_layout.addWidget(contador_preview)
        self.preview_layout.addWidget(cabecalho_preview)

        for indice, campo in enumerate(self.modelo_atual_campos):
            tipo = campo.get("tipo")
            label = campo.get("label")

            # --- Linha externa: conteúdo do campo (esquerda) + ferramentas (direita) ---
            linha_container = QFrame()
            linha_container.setStyleSheet("QFrame { border: 1px dashed #e2e8f0; border-radius: 6px; }")
            linha_layout = QHBoxLayout(linha_container)
            linha_layout.setContentsMargins(8, 6, 8, 6)
            linha_layout.setSpacing(8)

            ordem = QLabel(str(indice + 1))
            ordem.setFixedSize(24, 24)
            ordem.setAlignment(Qt.AlignmentFlag.AlignCenter)
            ordem.setStyleSheet(
                "QLabel { background: #e0f2fe; color: #0369a1; border-radius: 12px; font-weight: bold; }"
            )
            linha_layout.addWidget(ordem, alignment=Qt.AlignmentFlag.AlignTop)

            conteudo_layout = QVBoxLayout()
            conteudo_layout.setSpacing(2)

            if tipo == "secao":
                secao_label = QLabel(label.upper())
                secao_label.setStyleSheet("color: #0284c7 !important; font-size: 14px; font-weight: bold; background-color: transparent;")
                secao_frame = QFrame()
                secao_frame.setStyleSheet("background-color: #cbd5e1 !important; max-height: 1px; border: none;")
                conteudo_layout.addWidget(secao_label)
                conteudo_layout.addWidget(secao_frame)

            elif tipo == "texto_curto":
                lbl = QLabel(label)
                lbl.setStyleSheet(estilo_label)
                inp = QLineEdit()
                inp.setPlaceholderText("Área de Visualização do Input Curto")
                inp.setReadOnly(True)
                inp.setStyleSheet(estilo_input_curto)
                conteudo_layout.addWidget(lbl)
                conteudo_layout.addWidget(inp)

            elif tipo == "texto_longo":
                lbl = QLabel(label)
                lbl.setStyleSheet(estilo_label)
                inp = QTextEdit()
                inp.setPlaceholderText("Área de Visualização do Input Longo")
                inp.setReadOnly(True)
                inp.setMinimumHeight(50)
                inp.setMaximumHeight(70)
                inp.setStyleSheet(estilo_input_longo)
                conteudo_layout.addWidget(lbl)
                conteudo_layout.addWidget(inp)

            elif tipo == "checkbox":
                chk = QCheckBox(label)
                chk.setEnabled(False)
                chk.setStyleSheet("QCheckBox { color: #0f172a !important; font-size: 13px; font-weight: 500; padding: 4px; background-color: transparent; }")
                conteudo_layout.addWidget(chk)

            elif tipo == "numero":
                unidade = campo.get("unidade", "")
                lbl = QLabel(f"{label} ({unidade})" if unidade else label)
                lbl.setStyleSheet(estilo_label)
                inp = QLineEdit()
                inp.setPlaceholderText(f"Área de Visualização — número{f' em {unidade}' if unidade else ''}")
                inp.setReadOnly(True)
                inp.setStyleSheet(estilo_input_curto)
                conteudo_layout.addWidget(lbl)
                conteudo_layout.addWidget(inp)

            elif tipo == "data":
                lbl = QLabel(label)
                lbl.setStyleSheet(estilo_label)
                inp = QDateEdit()
                inp.setCalendarPopup(True)
                inp.setDate(QDate.currentDate())
                inp.setEnabled(False)
                inp.setStyleSheet("QDateEdit { background-color: #f8fafc !important; color: #0f172a !important; border: 1px solid #cbd5e1 !important; border-radius: 6px; padding: 6px 10px; font-size: 13px; }")
                conteudo_layout.addWidget(lbl)
                conteudo_layout.addWidget(inp)

            elif tipo == "multipla_escolha":
                lbl = QLabel(label)
                lbl.setStyleSheet(estilo_label)
                conteudo_layout.addWidget(lbl)
                for opcao in campo.get("opcoes", []):
                    radio = QRadioButton(opcao)
                    radio.setEnabled(False)
                    radio.setStyleSheet("QRadioButton { color: #0f172a !important; font-size: 13px; padding: 2px; background-color: transparent; }")
                    conteudo_layout.addWidget(radio)

            linha_layout.addLayout(conteudo_layout, stretch=1)

            # --- Ferramentas: mover pra cima/baixo, editar rótulo, excluir ---
            ferramentas_layout = QVBoxLayout()
            ferramentas_layout.setSpacing(2)

            linha_botoes_topo = QHBoxLayout()
            linha_botoes_topo.setSpacing(2)

            btn_subir = QPushButton("▲")
            btn_subir.setFixedSize(22, 20)
            btn_subir.setStyleSheet(estilo_btn_ferramenta)
            btn_subir.setEnabled(indice > 0)
            btn_subir.clicked.connect(lambda _, i=indice: self.mover_campo_rascunho(i, -1))
            linha_botoes_topo.addWidget(btn_subir)

            btn_descer = QPushButton("▼")
            btn_descer.setFixedSize(22, 20)
            btn_descer.setStyleSheet(estilo_btn_ferramenta)
            btn_descer.setEnabled(indice < total_campos - 1)
            btn_descer.clicked.connect(lambda _, i=indice: self.mover_campo_rascunho(i, 1))
            linha_botoes_topo.addWidget(btn_descer)

            ferramentas_layout.addLayout(linha_botoes_topo)

            if tipo == "secao":
                btn_converter = QPushButton("Campo")
                btn_converter.setToolTip("Transformar esta secao em campo de texto")
                btn_converter.setFixedHeight(20)
                btn_converter.setStyleSheet(estilo_btn_ferramenta)
                btn_converter.clicked.connect(
                    lambda _, i=indice: self.converter_secao_em_campo(i)
                )
                ferramentas_layout.addWidget(btn_converter)

            linha_botoes_baixo = QHBoxLayout()
            linha_botoes_baixo.setSpacing(2)

            btn_editar = QPushButton("✏️")
            btn_editar.setFixedSize(22, 20)
            btn_editar.setStyleSheet(estilo_btn_ferramenta)
            btn_editar.clicked.connect(lambda _, i=indice: self.editar_campo_rascunho(i))
            linha_botoes_baixo.addWidget(btn_editar)

            btn_excluir = QPushButton("🗑️")
            btn_excluir.setFixedSize(22, 20)
            btn_excluir.setStyleSheet("QPushButton { background-color: #fee2e2; color: #b91c1c; border: 1px solid #fca5a5; border-radius: 4px; font-size: 10px; padding: 2px; } QPushButton:hover { background-color: #fecaca; }")
            btn_excluir.clicked.connect(lambda _, i=indice: self.remover_campo_rascunho(i))
            linha_botoes_baixo.addWidget(btn_excluir)

            ferramentas_layout.addLayout(linha_botoes_baixo)

            linha_layout.addLayout(ferramentas_layout)

            self.preview_layout.addWidget(linha_container)

        self.scroll_content.adjustSize()

    def mover_campo_rascunho(self, indice, direcao):
        """Troca o campo de posição com o vizinho (direcao = -1 sobe, +1 desce)."""
        novo_indice = indice + direcao
        if 0 <= novo_indice < len(self.modelo_atual_campos):
            self.modelo_atual_campos[indice], self.modelo_atual_campos[novo_indice] = \
                self.modelo_atual_campos[novo_indice], self.modelo_atual_campos[indice]
            self.atualizar_visualizacao_preview()

    def converter_secao_em_campo(self, indice):
        """Permite corrigir uma secao importada que deveria ser preenchivel."""
        if not (0 <= indice < len(self.modelo_atual_campos)):
            return
        campo = self.modelo_atual_campos[indice]
        if campo.get("tipo") != "secao":
            return
        texto_sem_acento = unicodedata.normalize(
            "NFKD", campo.get("label", "campo")
        ).encode("ascii", "ignore").decode("ascii")
        campo["tipo"] = "texto_longo"
        campo["id"] = f"importado_{indice}_{re.sub(r'[^a-z0-9]+', '_', texto_sem_acento.lower()).strip('_')}"
        self.atualizar_visualizacao_preview()

    def editar_campo_rascunho(self, indice):
        """Abre um diálogo pré-preenchido para renomear o rótulo do campo/seção."""
        if not (0 <= indice < len(self.modelo_atual_campos)):
            return
        campo = self.modelo_atual_campos[indice]
        dial = CustomInputDialog("Editar Campo", "Novo texto para este campo/seção:", self)
        dial.input_field.setText(campo.get("label", ""))
        if dial.exec() == QDialog.Accepted:
            novo_texto = dial.get_text()
            if novo_texto:
                campo["label"] = novo_texto
                self.atualizar_visualizacao_preview()

    def remover_campo_rascunho(self, indice):
        if 0 <= indice < len(self.modelo_atual_campos):
            self.modelo_atual_campos.pop(indice)
            self.atualizar_visualizacao_preview()

    def salvar_modelo_customizado_db(self):
        if not self.modelo_atual_campos:
            self.exibir_popup("aviso", "Modelo Vazio", "Adicione pelo menos um campo antes de salvar.")
            return
            
        dial = CustomInputDialog("Salvar Modelo", "Dê um nome para este novo modelo de ficha:", self)
        if self._nome_modelo_importado and self._nome_modelo_importado != "Novo modelo":
            sugestao = os.path.splitext(os.path.basename(self._nome_modelo_importado))[0]
            dial.input_field.setText(sugestao)
        if dial.exec() == QDialog.Accepted:
            nome_modelo = dial.get_text()
            if not nome_modelo: return
        else:
            return
            
        estrutura_json = json.dumps(self.modelo_atual_campos, ensure_ascii=False)
        
        if not self.db.supabase:
            return
            
        try:
            payload = {
                "consultorio_id": self.db.consultorio_id,
                "nome_modelo": nome_modelo,
                "estrutura_json": estrutura_json
            }
            
            # Executa o upsert para evitar duplicações de modelos dentro do mesmo consultório
            self.db.supabase.table("modelos_fichas").upsert(
                payload, 
                on_conflict="consultorio_id,nome_modelo"
            ).execute()
            
            self.modo_criacao = False
            self._nome_modelo_importado = None
            self.carregar_modelos_iniciais_combo()
            self.combo_modelo.setCurrentText(nome_modelo)
            self.renderizar_formulario_dinamico()
            
            self.exibir_popup("info", "Sucesso!", f"O modelo '{nome_modelo}' foi gravado no banco de dados e já está pronto para uso!")
        except Exception as e:
            self.exibir_popup("erro", "Erro ao Salvar", f"Falha ao gravar no banco:\n{str(e)}")

    def _detectar_campos_a_partir_de_linhas(self, linhas_texto):
        """Heurística compartilhada: recebe uma lista de linhas de texto (extraídas
        de um .docx ou .pdf) e tenta identificar seções, campos de texto curto/longo
        e caixas de seleção. Usada tanto pelo importador de Word quanto de PDF."""
        novos_campos = []
        chaves_existentes = set()

        for texto in linhas_texto:
            if ":" in texto:
                partes = texto.split(":")
                label_campo = partes[0].replace(",", "").replace("-", "").strip()

                if not label_campo or len(label_campo) < 2 or len(label_campo) > 60:
                    continue

                id_campo = "".join(c for c in label_campo if c.isalnum()).lower()
                if id_campo in chaves_existentes:
                    continue
                chaves_existentes.add(id_campo)

                if "[" in texto or "]" in texto or "( )" in texto:
                    novos_campos.append({"tipo": "checkbox", "label": label_campo, "id": id_campo})
                else:
                    id_min = id_campo.lower()
                    if any(x in id_min for x in ["qp", "hda", "conduta", "antecedentes", "historico", "outros", "medicamentos", "observacoes"]):
                        tipo_campo = "texto_longo"
                    else:
                        tipo_campo = "texto_curto"
                    novos_campos.append({"tipo": tipo_campo, "label": label_campo, "id": id_campo})

            elif len(texto) < 50 and (texto.isupper() or len(texto) < 30):
                texto_limpo = texto.replace("-", "").replace(",", "").strip()
                if texto_limpo and len(texto_limpo) > 2:
                    novos_campos.append({"tipo": "secao", "label": texto_limpo})

        return novos_campos

    def _detectar_campos_melhorado(self, linhas_texto):
        """Interpreta estruturas frequentes de fichas Word/PDF para revisão no construtor."""
        campos = []
        ids = set()
        for linha in linhas_texto:
            texto = re.sub(r"\s+", " ", (linha or "")).strip(" -\t")
            if not texto or len(texto) > 180:
                continue

            opcoes = re.findall(r"\(\s*\)\s*([^()]{1,35}?)(?=\s*\(\s*\)|$)", texto)
            tem_checkbox = any(marcador in texto for marcador in ("[ ]", "☐", "□"))
            encontrado = re.match(r"^(.{2,80}?)(?:\s*:\s*|\s*[-–]\s+|\s*[_.]{3,}\s*$)", texto)

            if encontrado:
                label = encontrado.group(1).strip(" -:;,. ")
                texto_sem_acento = unicodedata.normalize("NFKD", label).encode("ascii", "ignore").decode("ascii")
                campo_id = re.sub(r"[^a-z0-9]+", "_", texto_sem_acento.lower()).strip("_")
                if not label or campo_id in ids:
                    continue
                ids.add(campo_id)
                palavras = campo_id.lower()
                if len(opcoes) >= 2:
                    campos.append({"tipo": "multipla_escolha", "label": label, "id": campo_id, "opcoes": opcoes})
                elif tem_checkbox:
                    campos.append({"tipo": "checkbox", "label": label, "id": campo_id})
                elif any(x in palavras for x in ("data", "nasc", "consulta", "atendimento")):
                    campos.append({"tipo": "data", "label": label, "id": campo_id})
                elif any(x in palavras for x in ("peso", "altura", "idade", "pressao", "temperatura", "frequencia", "glicemia")):
                    campos.append({"tipo": "numero", "label": label, "id": campo_id})
                elif any(x in palavras for x in ("qp", "hda", "conduta", "antecedentes", "historico", "medicamentos", "observacoes", "evolucao", "anamnese")):
                    campos.append({"tipo": "texto_longo", "label": label, "id": campo_id})
                else:
                    campos.append({"tipo": "texto_curto", "label": label, "id": campo_id})
            elif len(texto) <= 70:
                # Na importacao, uma linha isolada e tratada como campo. Isso
                # evita perder QP, HDA e outros rotulos clinicos em maiusculas.
                label = texto.rstrip(":").strip()
                texto_sem_acento = unicodedata.normalize("NFKD", label).encode("ascii", "ignore").decode("ascii")
                campo_id = re.sub(r"[^a-z0-9]+", "_", texto_sem_acento.lower()).strip("_")
                if not label or campo_id in ids:
                    continue
                ids.add(campo_id)
                palavras = campo_id.lower()
                tipo = "texto_longo" if any(
                    x in palavras for x in ("qp", "hda", "conduta", "antecedentes", "historico", "medicamentos", "observacoes", "evolucao", "anamnese", "exame", "avaliacao")
                ) else "texto_curto"
                campos.append({"tipo": tipo, "label": label, "id": campo_id})
        return campos

    def _aplicar_modelo_importado(self, novos_campos, nome_arquivo):
        if novos_campos:
            nome_reduzido = os.path.basename(nome_arquivo)
            self.iniciar_criacao_modelo(novos_campos, origem=nome_reduzido)
            self.exibir_popup(
                "info", "Documento interpretado",
                f"Encontramos {len(novos_campos)} elementos. Revise, edite e salve o modelo quando estiver satisfeito.",
            )
        else:
            self.exibir_popup("aviso", "Aviso", "Nenhum campo estruturado foi identificado.")

    def importar_modelo_word(self):
        if not DOCX_DISPONIVEL or Document is None:
            self.exibir_popup("erro", "Módulo Ausente", "A biblioteca 'python-docx' é necessária.\npip install python-docx")
            return

        file_path, _ = QFileDialog.getOpenFileName(self, "Selecionar Ficha em Word", "", "Arquivos Word (*.docx)")
        if not file_path: return
            
        try:
            doc = Document(file_path)
            linhas_texto = []
            
            for p in doc.paragraphs:
                t = p.text.strip()
                if t: linhas_texto.append(t)
            
            for tabela in doc.tables:
                for row_tab in tabela.rows:
                    valores_celulas = [c.text.strip() for c in row_tab.cells if c.text.strip()]
                    for val in valores_celulas:
                        for sub_val in val.split('\n'):
                            if sub_val.strip() and sub_val.strip() not in linhas_texto:
                                rastro = sub_val.strip()
                                if rastro.replace(",", "").strip():
                                    linhas_texto.append(rastro)

            if not linhas_texto:
                self.exibir_popup("aviso", "Documento sem texto", "Não encontramos texto no Word selecionado.")
                return
            novos_campos = self._detectar_campos_melhorado(linhas_texto)
            self._aplicar_modelo_importado(novos_campos, file_path)
        except Exception as e:
            self.exibir_popup("erro", "Falha de Leitura", f"Erro:\n{str(e)}")

    def importar_modelo_pdf(self):
        """Lê um PDF (anamnese, relatório, exame estruturado) e tenta montar
        um novo modelo de ficha automaticamente, com a mesma heurística usada
        para arquivos Word."""
        if not PDF_DISPONIVEL or PdfReader is None:
            self.exibir_popup("erro", "Módulo Ausente", "A biblioteca 'pypdf' é necessária.\npip install pypdf")
            return

        file_path, _ = QFileDialog.getOpenFileName(self, "Selecionar Ficha em PDF", "", "Arquivos PDF (*.pdf)")
        if not file_path:
            return

        try:
            leitor = PdfReader(file_path)
            linhas_texto = []
            for pagina in leitor.pages:
                texto_pagina = pagina.extract_text() or ""
                for linha in texto_pagina.split("\n"):
                    linha_limpa = linha.strip()
                    if linha_limpa:
                        linhas_texto.append(linha_limpa)

            if not linhas_texto:
                self.exibir_popup(
                    "aviso", "PDF sem texto",
                    "Este PDF parece ser uma imagem digitalizada. Use um PDF com texto selecionável ou monte o modelo manualmente.",
                )
                return
            novos_campos = self._detectar_campos_melhorado(linhas_texto)
            self._aplicar_modelo_importado(novos_campos, file_path)
        except Exception as e:
            self.exibir_popup("erro", "Falha de Leitura", f"Erro ao ler o PDF:\n{str(e)}")

    # ============================================================
    # ANEXOS (fotos / PDFs vinculados à ficha preenchida atual)
    # ============================================================
    def anexar_arquivos(self):
        caminhos, _ = QFileDialog.getOpenFileNames(self, "Selecionar Foto(s) ou PDF(s)", "", EXTENSOES_ANEXO_ACEITAS)
        if not caminhos:
            return
        for caminho in caminhos:
            self.arquivos_anexados.append(caminho)
        self.renderizar_anexos_thumbnails()

    def remover_anexo(self, indice):
        if 0 <= indice < len(self.arquivos_anexados):
            self.arquivos_anexados.pop(indice)
        self.renderizar_anexos_thumbnails()

    def visualizar_anexo_local(self, caminho):
        """Abre uma foto ampliada dentro do app, ou um PDF no visualizador padrão do sistema."""
        extensao = os.path.splitext(caminho)[1].lower()
        if extensao == ".pdf":
            QDesktopServices.openUrl(QUrl.fromLocalFile(caminho))
            return

        dialog = QDialog(self)
        dialog.setWindowTitle(os.path.basename(caminho))
        dialog.setStyleSheet("background-color: #ffffff;")
        layout_dialog = QVBoxLayout(dialog)
        lbl_imagem = QLabel()
        pixmap = QPixmap(caminho)
        if not pixmap.isNull():
            pixmap = pixmap.scaled(700, 700, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            lbl_imagem.setPixmap(pixmap)
        layout_dialog.addWidget(lbl_imagem)
        dialog.exec()

    def renderizar_anexos_thumbnails(self):
        # Limpa a tira de miniaturas
        while self.anexos_strip_layout.count():
            item = self.anexos_strip_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        for indice, caminho in enumerate(self.arquivos_anexados):
            extensao = os.path.splitext(caminho)[1].lower()
            miniatura = QFrame()
            miniatura.setFixedSize(72, 72)
            miniatura.setStyleSheet("QFrame { background-color: white; border: 1px solid #cbd5e1; border-radius: 6px; }")
            lay_mini = QVBoxLayout(miniatura)
            lay_mini.setContentsMargins(2, 2, 2, 2)
            lay_mini.setSpacing(0)

            if extensao == ".pdf":
                lbl = QLabel("📄\nPDF")
                lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                lbl.setStyleSheet("font-size: 11px; font-weight: bold; color: #0f172a; border: none; background: transparent;")
            else:
                lbl = QLabel()
                pixmap = QPixmap(caminho)
                if not pixmap.isNull():
                    pixmap = pixmap.scaled(64, 50, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                    lbl.setPixmap(pixmap)
                lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                lbl.setStyleSheet("border: none; background: transparent;")
            lay_mini.addWidget(lbl)

            btn_remover = QPushButton("✕")
            btn_remover.setFixedSize(16, 16)
            btn_remover.setStyleSheet("QPushButton { background-color: #fee2e2; color: #b91c1c; font-size: 10px; font-weight: bold; border: none; border-radius: 8px; } QPushButton:hover { background-color: #fca5a5; }")
            btn_remover.clicked.connect(lambda _, i=indice: self.remover_anexo(i))
            lay_mini.addWidget(btn_remover, alignment=Qt.AlignmentFlag.AlignRight)

            miniatura.mousePressEvent = lambda event, c=caminho: self.visualizar_anexo_local(c)

            self.anexos_strip_layout.addWidget(miniatura)

        self.anexos_strip_layout.addStretch()

    def _fazer_upload_anexos(self, paciente_id):
        """Envia cada arquivo pendente para o Supabase Storage e retorna a lista
        de metadados (nome, caminho no bucket, tipo) já pronta para salvar no banco."""
        lista_metadados = []
        if not self.arquivos_anexados or not self.db.supabase:
            return lista_metadados

        for caminho_local in self.arquivos_anexados:
            try:
                nome_original = os.path.basename(caminho_local)
                extensao = os.path.splitext(nome_original)[1]
                nome_unico = f"{uuid.uuid4().hex}{extensao}"
                caminho_no_bucket = f"{self.db.consultorio_id}/{paciente_id}/{nome_unico}"
                tipo_mime = mimetypes.guess_type(nome_original)[0] or "application/octet-stream"

                with open(caminho_local, "rb") as arquivo:
                    self.db.supabase.storage.from_(NOME_BUCKET_ANEXOS).upload(
                        caminho_no_bucket,
                        arquivo.read(),
                        {"content-type": tipo_mime}
                    )

                lista_metadados.append({
                    "nome": nome_original,
                    "caminho": caminho_no_bucket,
                    "tipo": tipo_mime
                })
            except Exception as e:
                print(f"Erro ao enviar anexo '{caminho_local}': {e}")

        return lista_metadados

    def iniciar_nova_ficha_para_paciente(self, paciente_id):
        """Prepara uma nova ficha para o paciente vindo da Agenda."""
        self.cancelar_edicao_ficha()
        self.carregar_pacientes_combo()
        indice = self.combo_paciente.findData(paciente_id)
        if indice >= 0:
            self.combo_paciente.setCurrentIndex(indice)
            self.lbl_status_operacao.setText("Nova ficha preparada para o paciente da consulta.")
            return True
        self.exibir_popup("erro", "Paciente não encontrado", "Não foi possível localizar o paciente da consulta.")
        return False

    def abrir_ficha_para_edicao(self, ficha_id):
        """Abre uma ficha existente no formulário original, já preenchida."""
        if not self.db.supabase:
            return False
        try:
            # O Supabase usa bigint para este campo; garante que a sessão não
            # deixe o identificador como texto ao restaurar dados locais.
            self.db.consultorio_id = int(self.db.consultorio_id)
            resposta = self.db.supabase.table("fichas_preenchidas").select(
                "id, paciente_id, modelo_nome, dados_respostas, data_atendimento, anexos"
            ).eq("id", ficha_id).eq("consultorio_id", self.db.consultorio_id).is_(
                "deleted_at", "null"
            ).execute()
            if not resposta.data:
                self.exibir_popup("erro", "Ficha não encontrada", "Essa ficha não está mais disponível para edição.")
                return False
            ficha = resposta.data[0]

            self.carregar_pacientes_combo()
            self.carregar_modelos_iniciais_combo()
            indice_paciente = self.combo_paciente.findData(ficha["paciente_id"])
            if indice_paciente >= 0:
                self.combo_paciente.setCurrentIndex(indice_paciente)

            try:
                respostas = ficha["dados_respostas"] if isinstance(ficha["dados_respostas"], dict) else json.loads(ficha["dados_respostas"] or "{}")
            except (TypeError, json.JSONDecodeError):
                respostas = {}

            nome_modelo = ficha.get("modelo_nome") or "Ficha de Consulta Geral (Padrão)"
            if "Padrão" in nome_modelo:
                self.combo_modelo.setCurrentText("Ficha de Consulta Geral (Padrão)")
                self.gerar_modelo_padrao()
            else:
                modelo = self.db.supabase.table("modelos_fichas").select("estrutura_json").eq(
                    "consultorio_id", self.db.consultorio_id
                ).eq("nome_modelo", nome_modelo).execute()
                if modelo.data:
                    estrutura = modelo.data[0]["estrutura_json"]
                    self.modelo_atual_campos = estrutura if isinstance(estrutura, list) else json.loads(estrutura)
                    self.combo_modelo.setCurrentText(nome_modelo)
                else:
                    self.modelo_atual_campos = self._campos_recuperados(respostas)
                self.renderizar_formulario_dinamico()

            self._preencher_formulario_com_respostas(respostas)
            self.ficha_em_edicao_id = ficha["id"]
            self._data_atendimento_original = ficha.get("data_atendimento")
            anexos = ficha.get("anexos") or []
            self._anexos_existentes = anexos if isinstance(anexos, list) else json.loads(anexos)
            self.arquivos_anexados = []
            self.combo_paciente.setEnabled(False)
            self.combo_modelo.setEnabled(False)
            self.btn_salvar_atendimento.setText("Salvar alterações da ficha")
            self.btn_cancelar_edicao.setVisible(True)
            self.lbl_status_operacao.setText("Editando ficha existente. O paciente e o modelo foram preservados.")
            return True
        except Exception as e:
            print(f"Erro ao abrir ficha para edição: {e}")
            self.exibir_popup("erro", "Não foi possível abrir", "Não foi possível preparar essa ficha para edição.")
            return False

    def _campos_recuperados(self, respostas):
        campos = []
        for campo_id, valor in respostas.items():
            tipo = "checkbox" if isinstance(valor, bool) else "texto_longo"
            label = campo_id.replace("custom_", "").replace("_", " ").capitalize()
            campos.append({"tipo": tipo, "label": label, "id": campo_id})
        return campos

    def _preencher_formulario_com_respostas(self, respostas):
        for campo_id, valor in respostas.items():
            if campo_id not in self.widgets_dinamicos:
                continue
            tipo, widget = self.widgets_dinamicos[campo_id]
            if tipo in ("texto_curto", "numero"):
                widget.setText("" if valor is None else str(valor))
            elif tipo == "texto_longo":
                widget.setPlainText("" if valor is None else str(valor))
            elif tipo == "checkbox":
                widget.setChecked(bool(valor))
            elif tipo == "data":
                data = QDate.fromString(str(valor), "dd/MM/yyyy")
                if data.isValid():
                    widget.setDate(data)
            elif tipo == "multipla_escolha":
                for botao in widget.buttons():
                    if botao.text() == str(valor):
                        botao.setChecked(True)
                        break

    def cancelar_edicao_ficha(self):
        self.ficha_em_edicao_id = None
        self._data_atendimento_original = None
        self._anexos_existentes = []
        self.combo_paciente.setEnabled(True)
        self.combo_modelo.setEnabled(True)
        self.btn_salvar_atendimento.setText("Salvar Ficha Preenchida")
        self.btn_cancelar_edicao.setVisible(False)
        self.lbl_status_operacao.setText("Edição cancelada. Você pode iniciar uma nova ficha.")
        self.gerar_modelo_padrao()

    def renderizar_formulario_dinamico(self):
        self.limpar_layout_completamente(self.dinamic_form_layout)
        self.widgets_dinamicos.clear()
        
        estilo_label = "color: #334155 !important; font-weight: bold; font-size: 13px; margin-top: 8px; background-color: transparent;"
        estilo_input_curto = """
            QLineEdit { background-color: #ffffff !important; color: #0f172a !important; border: 1px solid #cbd5e1 !important; border-radius: 6px; padding: 8px 12px; font-size: 13px; }
            QLineEdit:focus { border: 1px solid #0284c7 !important; }
        """
        estilo_input_longo = """
            QTextEdit { background-color: #ffffff !important; color: #0f172a !important; border: 1px solid #cbd5e1 !important; border-radius: 6px; padding: 8px 12px; font-size: 13px; }
            QTextEdit:focus { border: 1px solid #0284c7 !important; }
        """
        
        for campo in self.modelo_atual_campos:
            tipo = campo.get("tipo")
            label = campo.get("label")
            id_campo = campo.get("id")
            
            if tipo == "secao":
                secao_frame = QFrame()
                secao_frame.setStyleSheet("background-color: #cbd5e1 !important; max-height: 1px; border: none; margin-top: 4px;")
                secao_label = QLabel(label.upper())
                secao_label.setStyleSheet("color: #0284c7 !important; font-size: 14px; font-weight: bold; margin-top: 18px; background-color: transparent;")
                self.dinamic_form_layout.addWidget(secao_label)
                self.dinamic_form_layout.addWidget(secao_frame)
                
            elif tipo == "texto_curto":
                lbl = QLabel(label)
                lbl.setStyleSheet(estilo_label)
                inp = QLineEdit()
                inp.setStyleSheet(estilo_input_curto)
                if "placeholder" in campo: inp.setPlaceholderText(campo["placeholder"])
                self.dinamic_form_layout.addWidget(lbl)
                self.dinamic_form_layout.addWidget(inp)
                self.widgets_dinamicos[id_campo] = ("texto_curto", inp)
                
            elif tipo == "texto_longo":
                lbl = QLabel(label)
                lbl.setStyleSheet(estilo_label)
                inp = QTextEdit()
                inp.setMinimumHeight(65)
                inp.setMaximumHeight(120)
                inp.setStyleSheet(estilo_input_longo)
                self.dinamic_form_layout.addWidget(lbl)
                self.dinamic_form_layout.addWidget(inp)
                self.widgets_dinamicos[id_campo] = ("texto_longo", inp)
                
            elif tipo == "checkbox":
                chk = QCheckBox(label)
                chk.setStyleSheet("QCheckBox { color: #0f172a !important; font-size: 13px; font-weight: 500; padding: 4px; background-color: transparent; }")
                self.dinamic_form_layout.addWidget(chk)
                self.widgets_dinamicos[id_campo] = ("checkbox", chk)

            elif tipo == "numero":
                unidade = campo.get("unidade", "")
                lbl = QLabel(f"{label} ({unidade})" if unidade else label)
                lbl.setStyleSheet(estilo_label)
                inp = QLineEdit()
                inp.setValidator(QDoubleValidator())
                inp.setPlaceholderText(f"Digite um número{f' em {unidade}' if unidade else ''}...")
                inp.setStyleSheet(estilo_input_curto)
                self.dinamic_form_layout.addWidget(lbl)
                self.dinamic_form_layout.addWidget(inp)
                self.widgets_dinamicos[id_campo] = ("numero", inp)

            elif tipo == "data":
                lbl = QLabel(label)
                lbl.setStyleSheet(estilo_label)
                inp = QDateEdit()
                inp.setCalendarPopup(True)
                inp.setDisplayFormat("dd/MM/yyyy")
                inp.setDate(QDate.currentDate())
                inp.setStyleSheet("QDateEdit { background-color: #ffffff !important; color: #0f172a !important; border: 1px solid #cbd5e1 !important; border-radius: 6px; padding: 7px 10px; font-size: 13px; }")
                self.dinamic_form_layout.addWidget(lbl)
                self.dinamic_form_layout.addWidget(inp)
                self.widgets_dinamicos[id_campo] = ("data", inp)

            elif tipo == "multipla_escolha":
                lbl = QLabel(label)
                lbl.setStyleSheet(estilo_label)
                self.dinamic_form_layout.addWidget(lbl)
                grupo_radio = QButtonGroup(self)
                for opcao in campo.get("opcoes", []):
                    radio = QRadioButton(opcao)
                    radio.setStyleSheet("QRadioButton { color: #0f172a !important; font-size: 13px; padding: 3px; background-color: transparent; }")
                    grupo_radio.addButton(radio)
                    self.dinamic_form_layout.addWidget(radio)
                self.widgets_dinamicos[id_campo] = ("multipla_escolha", grupo_radio)

        self.dinamic_form_layout.addStretch()
        self.scroll_content.adjustSize()

    def _dados_para_exportacao(self):
        paciente = self.combo_paciente.currentText().strip()
        if not paciente or not self.combo_paciente.currentData():
            self.exibir_popup("aviso", "Selecione o paciente", "Selecione um paciente antes de exportar a ficha.")
            return None

        rotulos = {campo.get("id"): campo.get("label", campo.get("id", "")) for campo in self.modelo_atual_campos}
        respostas = []
        for campo_id, (tipo, widget) in self.widgets_dinamicos.items():
            if tipo in ("texto_curto", "numero"):
                valor = widget.text().strip()
            elif tipo == "texto_longo":
                valor = widget.toPlainText().strip()
            elif tipo == "checkbox":
                valor = "Sim" if widget.isChecked() else "Não"
            elif tipo == "data":
                valor = widget.date().toString("dd/MM/yyyy")
            elif tipo == "multipla_escolha":
                marcado = widget.checkedButton()
                valor = marcado.text() if marcado else ""
            else:
                valor = ""
            if valor:
                respostas.append((rotulos.get(campo_id, campo_id.replace("_", " ").capitalize()), valor))

        profissional = self.db.obter_nome_profissional() if hasattr(self.db, "obter_nome_profissional") else ""
        clinica = getattr(getattr(self.db, "session_manager", None), "nome_clinica", None) or "Prontu"
        return {
            "paciente": paciente,
            "modelo": self.combo_modelo.currentText() or "Ficha Clínica",
            "data": datetime.now().strftime("%d/%m/%Y às %H:%M"),
            "clinica": clinica,
            "profissional": profissional,
            "respostas": respostas,
        }

    @staticmethod
    def _nome_arquivo_exportacao(dados, extensao):
        nome = re.sub(r"[^A-Za-z0-9_-]+", "_", dados["paciente"].strip())
        return f"Ficha_{nome}_{datetime.now().strftime('%Y%m%d_%H%M')}{extensao}"

    def exportar_ficha_word(self):
        dados = self._dados_para_exportacao()
        if not dados:
            return
        if not DOCX_DISPONIVEL or Document is None:
            self.exibir_popup("erro", "Word indisponível", "A biblioteca python-docx não está instalada.")
            return
        sugestao = self._nome_arquivo_exportacao(dados, ".docx")
        caminho, _ = QFileDialog.getSaveFileName(self, "Salvar ficha em Word", sugestao, "Documento Word (*.docx)")
        if not caminho:
            return
        try:
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            documento = Document()
            secao = documento.sections[0]
            secao.top_margin = Inches(0.65)
            secao.bottom_margin = Inches(0.65)
            cabecalho = documento.add_table(rows=1, cols=2)
            cabecalho.columns[0].width = Inches(0.8)
            celula_logo, celula_texto = cabecalho.rows[0].cells
            logo = os.path.join(os.path.dirname(__file__), "..", "assets", "prontu_logo.png")
            if os.path.exists(logo):
                celula_logo.paragraphs[0].add_run().add_picture(logo, width=Inches(0.48))
            paragrafo = celula_texto.paragraphs[0]
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragrafo.add_run(dados["clinica"])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(2, 132, 199)

            titulo = documento.add_heading(dados["modelo"], level=0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            documento.add_paragraph(f"Paciente: {dados['paciente']}")
            documento.add_paragraph(f"Data de emissão: {dados['data']}")
            if dados["profissional"]:
                documento.add_paragraph(f"Profissional: {dados['profissional']}")
            documento.add_paragraph("─" * 55)
            for rotulo, valor in dados["respostas"]:
                documento.add_heading(rotulo, level=2)
                documento.add_paragraph(valor)
            documento.add_paragraph("\nDocumento gerado pelo Prontu.")
            documento.save(caminho)
            self.lbl_status_operacao.setText("Ficha exportada em Word com sucesso.")
        except Exception as e:
            self.exibir_popup("erro", "Erro ao exportar", f"Não foi possível criar o Word.\n{e}")

    def exportar_ficha_pdf(self):
        dados = self._dados_para_exportacao()
        if not dados:
            return
        sugestao = self._nome_arquivo_exportacao(dados, ".pdf")
        caminho, _ = QFileDialog.getSaveFileName(self, "Salvar ficha em PDF", sugestao, "Documento PDF (*.pdf)")
        if not caminho:
            return
        try:
            campos_html = "".join(
                f"<section><h2>{html.escape(rotulo)}</h2><p>{html.escape(str(valor)).replace(chr(10), '<br>')}</p></section>"
                for rotulo, valor in dados["respostas"]
            ) or "<p>Nenhuma resposta preenchida.</p>"
            profissional_html = f"<p><strong>Profissional:</strong> {html.escape(dados['profissional'])}</p>" if dados["profissional"] else ""
            conteudo = f"""
                <html><head><style>
                    body {{ font-family: Arial; color: #1e293b; font-size: 10.5pt; }}
                    .cabecalho {{ border-bottom: 2px solid #0284c7; padding-bottom: 12px; }}
                    h1 {{ color: #0f172a; font-size: 19pt; margin: 18px 0 8px; }}
                    h2 {{ color: #0284c7; font-size: 12pt; margin: 18px 0 5px; }}
                    p {{ line-height: 1.4; margin: 4px 0; }}
                    section {{ border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
                    .rodape {{ color: #64748b; font-size: 8pt; margin-top: 24px; }}
                </style></head><body>
                    <div class='cabecalho'><strong>{html.escape(dados['clinica'])}</strong></div>
                    <h1>{html.escape(dados['modelo'])}</h1>
                    <p><strong>Paciente:</strong> {html.escape(dados['paciente'])}</p>
                    <p><strong>Data de emissão:</strong> {html.escape(dados['data'])}</p>
                    {profissional_html}
                    {campos_html}
                    <p class='rodape'>Documento gerado pelo Prontu.</p>
                </body></html>
            """
            documento = QTextDocument()
            documento.setHtml(conteudo)
            impressora = QPrinter(QPrinter.PrinterMode.HighResolution)
            impressora.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            impressora.setOutputFileName(caminho)
            documento.print_(impressora)
            self.lbl_status_operacao.setText("Ficha exportada em PDF com sucesso.")
        except Exception as e:
            self.exibir_popup("erro", "Erro ao exportar", f"Não foi possível criar o PDF.\n{e}")

    def salvar_ficha_preenchida(self):
        if self.modo_criacao:
            self.exibir_popup("aviso", "Modo Construtor", "Você está editando um modelo de ficha. Clique em 'Concluir' primeiro.")
            return
            
        paciente_id = self.combo_paciente.currentData()
        if not paciente_id:
            self.exibir_popup("aviso", "Erro", "Selecione um paciente antes de salvar a ficha.")
            return
        
        respostas = {}
        for id_campo, (tipo, widget) in self.widgets_dinamicos.items():
            if tipo == "texto_curto": respostas[id_campo] = widget.text().strip()
            elif tipo == "texto_longo": respostas[id_campo] = widget.toPlainText().strip()
            elif tipo == "checkbox": respostas[id_campo] = widget.isChecked()
            elif tipo == "numero": respostas[id_campo] = widget.text().strip()
            elif tipo == "data": respostas[id_campo] = widget.date().toString("dd/MM/yyyy")
            elif tipo == "multipla_escolha":
                botao_marcado = widget.checkedButton()
                respostas[id_campo] = botao_marcado.text() if botao_marcado else ""
                
        modelo_nome = self.combo_modelo.currentText()
        data_atual = self._data_atendimento_original or datetime.now().strftime("%d/%m/%Y %H:%M")
        
        if not self.db.supabase:
            return
            
        try:
            # A API recebe estes campos como bigint no Supabase. A conversão
            # aqui evita que IDs restaurados como texto quebrem o salvamento.
            consultorio_id = int(self.db.consultorio_id)
            paciente_id = int(paciente_id)
            ficha_em_edicao_id = int(self.ficha_em_edicao_id) if self.ficha_em_edicao_id is not None else None
            editando_ficha = ficha_em_edicao_id is not None
            if editando_ficha:
                # Edição não altera os vínculos da ficha: atualiza somente as respostas.
                self.db.supabase.table("fichas_preenchidas").update(
                    {"dados_respostas": respostas}
                ).eq("id", ficha_em_edicao_id).execute()
                ficha_id = ficha_em_edicao_id
            else:
                payload = {
                    "consultorio_id": consultorio_id,
                    "paciente_id": paciente_id,
                    "modelo_nome": modelo_nome,
                    "dados_respostas": respostas,
                    "data_atendimento": data_atual,
                }
                resposta_insert = self.db.supabase.table("fichas_preenchidas").insert(payload).execute()
                ficha_id = resposta_insert.data[0]["id"] if resposta_insert.data else None

            # Se houver fotos/PDFs pendentes, envia agora para o Storage e vincula à ficha recém-criada
            if self.arquivos_anexados and ficha_id:
                metadados_anexos = self._fazer_upload_anexos(paciente_id)
                if metadados_anexos:
                    anexos_finais = self._anexos_existentes + metadados_anexos
                    self.db.supabase.table("fichas_preenchidas")\
                        .update({"anexos": json.dumps(anexos_finais, ensure_ascii=False)})\
                        .eq("id", ficha_id)\
                        .execute()
            
            # Limpa os campos após salvar com sucesso
            for tipo, widget in self.widgets_dinamicos.values():
                if tipo == "texto_curto": widget.clear()
                elif tipo == "texto_longo": widget.clear()
                elif tipo == "checkbox": widget.setChecked(False)
                elif tipo == "numero": widget.clear()
                elif tipo == "data": widget.setDate(QDate.currentDate())
                elif tipo == "multipla_escolha": widget.setExclusive(False); [b.setChecked(False) for b in widget.buttons()]; widget.setExclusive(True)

            self.arquivos_anexados = []
            self.renderizar_anexos_thumbnails()
            self._formulario_sujo = False
                
            if editando_ficha:
                self.cancelar_edicao_ficha()
                janela = getattr(self, "window_principal", None)
                tela_pacientes = getattr(janela, "screen_pacientes", None) if janela else None
                if tela_pacientes and tela_pacientes.id_em_edicao == paciente_id:
                    tela_pacientes.carregar_historico_fichas_paciente(paciente_id)
                self.exibir_popup("info", "Ficha atualizada", "As alterações foram salvas no atendimento original.")
            else:
                self.exibir_popup("info", "Ficha Salva", "O atendimento foi registrado com sucesso!")
        except Exception as e:
            self.exibir_popup("erro", "Erro", f"Falha no banco de dados:\n{str(e)}")
